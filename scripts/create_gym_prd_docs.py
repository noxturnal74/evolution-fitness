from pathlib import Path
from datetime import date
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "prd-documents-platform-v1.2"
PRESET_NAME = "standard_business_brief"
HEADER_PATTERN = "memo_masthead"
DOC_DATE = "May 25, 2026"


def rgb(value):
    value = value.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


BLACK = rgb("#000000")
MUTED = rgb("#555555")
BLUE = rgb("#2E74B5")
BLUE_DARK = rgb("#1F4D78")
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
TABLE_BORDER = "D9E2EC"


ARCHETYPES = {
    "local-community": {
        "label": "Local community gym",
        "tone": "Friendly, direct, energetic, and practical.",
        "visual": "Local city identity, energetic training photos, clear offers, and visible WhatsApp/Maps actions.",
        "journey": [
            "Hero with location and immediate WhatsApp CTA",
            "Why train here / trust points",
            "Facilities and training area proof",
            "Classes or programs",
            "Membership and visit options",
            "Personal training or community proof",
            "Location, hours, Maps, Instagram",
            "FAQ and final join CTA",
        ],
        "palette": "Brand accent from logo plus black, white, warm neutral, and one support accent.",
        "typography": "Readable modern sans-serif with bold compact display heading.",
    },
    "chain-modern": {
        "label": "Modern chain / network gym",
        "tone": "Confident, structured, modern, and conversion-oriented.",
        "visual": "Clean corporate layout, bento facilities, membership comparison, app-like interaction, and strong network credibility.",
        "journey": [
            "Hero with brand scale and primary join CTA",
            "Network/access proposition",
            "Facility ecosystem bento grid",
            "Classes, Pilates, or program universe",
            "Membership comparison",
            "Club/location finder style section",
            "Proof, FAQ, and final join flow",
        ],
        "palette": "Logo-led primary color, clean dark/light surfaces, restrained gradients, and high contrast CTA.",
        "typography": "Modern geometric sans-serif with strong numeric hierarchy.",
    },
    "hardcore": {
        "label": "Hardcore strength / performance gym",
        "tone": "Bold, focused, disciplined, and no-nonsense.",
        "visual": "Industrial dark surfaces, gritty equipment imagery, strong contrast, and challenge-based conversion.",
        "journey": [
            "Hero with performance promise",
            "Training culture and discipline section",
            "Equipment and free-weight proof",
            "Strength/PT programs",
            "Transformation or challenge gallery",
            "Membership and visit options",
            "Location and direct CTA",
            "FAQ and final challenge CTA",
        ],
        "palette": "Black/charcoal base, brand red/orange/gold/green accent, concrete gray, and white typography.",
        "typography": "Condensed athletic display heading plus readable sans-serif body.",
    },
    "women-wellness": {
        "label": "Women-friendly / wellness fitness",
        "tone": "Clean, confident, calm, supportive, and privacy-aware.",
        "visual": "Soft premium palette, bright studio imagery, comfort cues, class/community proof, and gentle motion.",
        "journey": [
            "Hero with comfort and confidence promise",
            "Privacy, support, and beginner-friendly reasons",
            "Facilities, pool, or studio highlights",
            "Class schedule and instructor support",
            "Membership or class pass options",
            "Gallery, testimonials, and community proof",
            "Location, Maps, and direct contact",
            "FAQ and final consultation CTA",
        ],
        "palette": "Soft neutral base with logo accent, warm light surfaces, and high contrast CTA.",
        "typography": "Elegant but readable sans-serif; avoid decorative serif and tiny body text.",
    },
    "cafe-lifestyle": {
        "label": "Fitness plus cafe / lifestyle",
        "tone": "Warm, modern, active, friendly, and social.",
        "visual": "Earthy warm palette, cafe/recovery lifestyle imagery, training photos, and community-first layout.",
        "journey": [
            "Hero with train plus recover proposition",
            "Gym and recovery concept",
            "Fitness packages",
            "Cafe, healthy food, or social space",
            "Classes and coaching",
            "Facilities and gallery",
            "Location, Maps, and reservation/contact",
            "FAQ and final CTA",
        ],
        "palette": "Warm neutral, espresso/cream support, one active fitness accent, and clean white cards.",
        "typography": "Friendly modern sans-serif with clear headings and relaxed spacing.",
    },
    "sport-complex": {
        "label": "Sport complex / multi-activity club",
        "tone": "Active, complete, family-friendly, and practical.",
        "visual": "Multi-zone layout, activity cards, clear facility map, and booking-friendly CTA.",
        "journey": [
            "Hero with complete sport facility proposition",
            "Activity ecosystem",
            "Facility zone tabs",
            "Membership and booking options",
            "Classes/coaching",
            "Gallery and proof",
            "Location, Maps, and contact",
            "FAQ and final CTA",
        ],
        "palette": "Sporty primary accent, neutral background, clear category colors, and accessible CTA contrast.",
        "typography": "Clean sans-serif with strong section labels and simple cards.",
    },
    "premium-wellness": {
        "label": "Premium hotel / wellness fitness",
        "tone": "Premium, calm, health-focused, and polished.",
        "visual": "Cinematic hotel/wellness cues, pool or studio imagery, refined spacing, and clear class/PT journey.",
        "journey": [
            "Hero with hotel/wellness positioning",
            "Facilities and wellness value",
            "Classes and training options",
            "Personal trainer journey",
            "Membership or visit options",
            "Gallery and social proof",
            "Location, Maps, and hours",
            "FAQ and final consultation CTA",
        ],
        "palette": "Deep green/blue/black base, clean white text, warm metallic or aqua accent.",
        "typography": "Premium modern sans-serif with generous spacing and readable body.",
    },
}


GYMS = [
    {
        "no": 1,
        "name": "Evolution Fitness Malang",
        "slug": "evolution-fitness-malang",
        "city": "Malang",
        "archetype": "local-community",
        "positioning": "Local Malang gym with strong class, community, promo, and direct WhatsApp conversion.",
        "primary_source": "https://latihanfisik.com/gym-detail/evolution-fitness-malang/",
        "instagram": "https://www.instagram.com/evolution_fitness_malang/",
        "whatsapp": "https://wa.me/6282228862049",
        "maps": "https://www.google.com/maps/search/?api=1&query=Evolution+Fitness+Malang",
        "logo_source": "https://www.instagram.com/evolution_fitness_malang/",
        "facts": [
            "Local Malang gym profile from LatihanFisik.",
            "Instagram should be used for logo, class visual, community proof, and promo direction.",
            "Primary contact CTA should route to WhatsApp 0822-2886-2049.",
        ],
        "hero": {
            "headline": "Evolution Fitness Malang",
            "subhead": "A local Malang training space built around classes, community energy, and direct support from the gym team.",
            "cta_primary": "Chat Evolution",
            "cta_secondary": "Open Maps",
        },
        "facility_focus": ["weights area", "studio classes", "personal training", "member community"],
        "program_focus": ["group classes", "promo campaigns", "PT consultation", "beginner guidance"],
    },
    {
        "no": 2,
        "name": "FTL Gym",
        "slug": "ftl-gym",
        "city": "Indonesia",
        "archetype": "chain-modern",
        "positioning": "Modern gym network with 24/7 access, club scale, group exercise, Pilates, and premium lifestyle facilities.",
        "primary_source": "https://ftlgym.com/ and https://ftlgym.com/membership/",
        "instagram": "https://www.instagram.com/ftlgym.id/",
        "whatsapp": "https://wa.me/62818687858",
        "maps": "https://www.google.com/maps/search/?api=1&query=FTL+Gym",
        "logo_source": "https://ftlgym.com/",
        "facts": [
            "Official website is the main source for membership and facilities.",
            "Instagram/source notes mention 60+ clubs, 24-hour access, Les Mills, Pilates, and premium facilities.",
            "WhatsApp CTA should use 0818-687-858.",
        ],
        "hero": {
            "headline": "FTL Gym",
            "subhead": "An all-in fitness ecosystem with club access, group classes, Pilates, and lifestyle facilities for everyday training.",
            "cta_primary": "Ask Membership",
            "cta_secondary": "Explore Clubs",
        },
        "facility_focus": ["24/7 access", "group exercise", "Pilates", "sauna/steam/jacuzzi or onsen where available"],
        "program_focus": ["membership", "Les Mills", "Pilates+", "club network"],
    },
    {
        "no": 3,
        "name": "Fitness Plus Dinoyo",
        "slug": "fitness-plus-dinoyo",
        "city": "Malang",
        "archetype": "chain-modern",
        "positioning": "Corporate mega gym in Dinoyo with 24-hour fitness value, broad equipment depth, classes, and PT.",
        "primary_source": "https://fitnessplusindonesia.co.id/",
        "instagram": "https://www.instagram.com/fitnessplus.dinoyo/",
        "whatsapp": "https://wa.me/6281130051777",
        "maps": "https://www.google.com/maps/search/?api=1&query=Fitness+Plus+Dinoyo+Malang",
        "logo_source": "https://fitnessplusindonesia.co.id/themes1/images/logo/logo.png",
        "facts": [
            "Official website and Dinoyo Instagram are primary data sources.",
            "Source notes mention 100+ international equipment, fun classes, and mega gym concept.",
            "Central WhatsApp CTA should use 0811-3005-1777.",
        ],
        "hero": {
            "headline": "Fitness Plus Dinoyo",
            "subhead": "A mega gym experience in Malang with broad equipment access, energetic classes, and a clear path to start training.",
            "cta_primary": "Ask Fitness Plus",
            "cta_secondary": "View Dinoyo Location",
        },
        "facility_focus": ["100+ international equipment", "class studio", "personal trainer", "24-hour mega gym positioning"],
        "program_focus": ["fun classes", "membership", "PT", "facility tour"],
    },
    {
        "no": 4,
        "name": "FITX Gym",
        "slug": "fitx-gym",
        "city": "Indonesia",
        "archetype": "chain-modern",
        "positioning": "Futuristic premium-affordable 24/7 gym with large equipment inventory, class scale, Pilates, and lifestyle services.",
        "primary_source": "https://fitxgym.id/ and https://linktr.ee/fitxgym.id",
        "instagram": "https://www.instagram.com/fitxgym.id/",
        "whatsapp": "https://wa.me/6281188021085",
        "maps": "https://www.google.com/maps/search/?api=1&query=FITX+Gym",
        "logo_source": "https://fitxgym.id/",
        "facts": [
            "Official website is main source for features.",
            "Source notes mention 24/7 access, 90+ machines, 1,000+ classes/month, sauna, ice bath, nail art, barbershop, and Pilates.",
            "Linktree can support free trial, membership, and branch CTA flow.",
        ],
        "hero": {
            "headline": "FITX Gym",
            "subhead": "A modern 24/7 training destination with machines, classes, recovery, and lifestyle services in one high-energy ecosystem.",
            "cta_primary": "Start With FITX",
            "cta_secondary": "View Facilities",
        },
        "facility_focus": ["90+ machines", "1,000+ classes/month", "sauna", "ice bath", "nail art", "barbershop", "Pilates"],
        "program_focus": ["24/7 membership", "free trial", "Pilates", "class schedule", "recovery"],
    },
    {
        "no": 5,
        "name": "Osbond Gym",
        "slug": "osbond-gym",
        "city": "Indonesia",
        "archetype": "chain-modern",
        "positioning": "National mega gym with one-stop body needs, all-club access, class depth, and premium facility proof.",
        "primary_source": "https://www.osbondgym.com/ and https://www.osbondgym.com/club",
        "instagram": "https://www.instagram.com/osbondgym/",
        "whatsapp": "https://www.osbondgym.com/contact-us",
        "maps": "https://www.google.com/maps/search/?api=1&query=Osbond+Gym",
        "logo_source": "https://www.osbondgym.com/imgproxy/images/navbar/osbond-logo-white.png",
        "facts": [
            "Official website and club page should drive content.",
            "Source notes mention 30+ club, premium equipment, swimming pool, sauna, steam, ice bath, and 160+ classes.",
            "Contact CTA should route to the official contact/free trial flow instead of inventing a phone number.",
        ],
        "hero": {
            "headline": "Osbond Gym",
            "subhead": "A one-stop fitness club network with classes, premium equipment, recovery facilities, and all-club access positioning.",
            "cta_primary": "Request Free Trial",
            "cta_secondary": "Explore Clubs",
        },
        "facility_focus": ["30+ club", "premium equipment", "swimming pool", "sauna", "steam", "ice bath"],
        "program_focus": ["160+ classes", "membership", "club locator", "free trial"],
    },
    {
        "no": 6,
        "name": "Belle Crown Gym",
        "slug": "belle-crown-gym",
        "city": "Malang",
        "archetype": "women-wellness",
        "positioning": "Wellness gym and pool in Malang with semi-premium comfort, 10+ classes, PT, and complete facilities.",
        "primary_source": "https://bellecrown.id/tempat-fitnes-gym-kolam-renang-di-malang.php and https://linktr.ee/bellecrownindonesia",
        "instagram": "https://www.instagram.com/bellecrown.gym/",
        "whatsapp": "https://linktr.ee/bellecrownindonesia",
        "maps": "https://www.google.com/maps/search/?api=1&query=Belle+Crown+Gym+Lembah+Dieng+Malang",
        "logo_source": "https://www.instagram.com/bellecrown.gym/",
        "facts": [
            "Official page and Linktree are sources for facility and CTA routing.",
            "Source notes mention Jl. Lembah Dieng 6, complete facilities, 10+ classes, swimming pool, and personal trainer.",
            "CTA should use Linktree if exact WA per branch is not confirmed.",
        ],
        "hero": {
            "headline": "Belle Crown Gym",
            "subhead": "A complete Malang wellness destination with gym access, swimming pool, classes, and personal trainer support.",
            "cta_primary": "Ask Belle Crown",
            "cta_secondary": "View Location",
        },
        "facility_focus": ["gym", "swimming pool", "10+ classes", "personal trainer", "complete facilities"],
        "program_focus": ["class schedule", "pool access", "PT", "membership/pricelist via Linktree"],
    },
    {
        "no": 7,
        "name": "Draco Gym",
        "slug": "draco-gym",
        "city": "Surabaya",
        "archetype": "hardcore",
        "positioning": "Dark-bold local gym in Kalijudan Surabaya with PT, direct WhatsApp contact, and practical opening hours.",
        "primary_source": "https://latihanfisik.com/gym-detail/draco-gym/",
        "instagram": "https://www.instagram.com/draco.gym_/",
        "whatsapp": "https://wa.me/62895385588866",
        "maps": "https://www.google.com/maps/search/?api=1&query=Draco+Gym+Kalijudan+Surabaya",
        "logo_source": "https://www.instagram.com/draco.gym_/",
        "facts": [
            "LatihanFisik and Instagram are main data sources.",
            "Source notes mention Kalijudan Surabaya, personal trainer, WhatsApp, and operating hours 06.00-22.00.",
            "Brand should feel bold and performance-led without inventing facility claims.",
        ],
        "hero": {
            "headline": "Draco Gym",
            "subhead": "A focused Surabaya training space for members who want direct support, strong workouts, and simple local access.",
            "cta_primary": "Chat Draco Gym",
            "cta_secondary": "Open Kalijudan Maps",
        },
        "facility_focus": ["strength area", "personal trainer", "local gym access", "daily training"],
        "program_focus": ["PT consultation", "strength program", "membership inquiry", "operating hours"],
    },
    {
        "no": 8,
        "name": "Planet Gym Surabaya",
        "slug": "planet-gym-surabaya",
        "city": "Surabaya",
        "archetype": "local-community",
        "positioning": "Accessible city gym and aerobic center on Banyu Urip with workout area, classes, PT, and parking.",
        "primary_source": "Instagram official",
        "instagram": "https://www.instagram.com/planetgymsby/",
        "whatsapp": "https://wa.me/628113451228",
        "maps": "https://www.google.com/maps/search/?api=1&query=Planet+Gym+Banyu+Urip+248+Surabaya",
        "logo_source": "https://www.instagram.com/planetgymsby/",
        "facts": [
            "Official Instagram is the main source.",
            "Source notes mention workout area, boxing/free-weight area, cardio, aerobic, group classes, personal training, parking, Banyu Urip 248, and contacts.",
            "Secondary contact can include telephone (031) 5467371.",
        ],
        "hero": {
            "headline": "Planet Gym Surabaya",
            "subhead": "A practical Surabaya fitness center with workout areas, aerobic classes, personal training, and easy Banyu Urip access.",
            "cta_primary": "Chat Planet Gym",
            "cta_secondary": "Open Maps",
        },
        "facility_focus": ["workout area", "boxing/free-weight area", "cardio", "aerobic", "parking"],
        "program_focus": ["group classes", "personal training", "aerobic", "membership inquiry"],
    },
    {
        "no": 9,
        "name": "New Icon Gym",
        "slug": "new-icon-gym",
        "city": "Surabaya",
        "archetype": "local-community",
        "positioning": "Local Surabaya gym at Barata Jaya with clear location, contact, gym floor, aerobic, trainer, and price discovery.",
        "primary_source": "https://latihanfisik.com/gym-detail/new-icon-gym/",
        "instagram": "https://www.instagram.com/new.icon.gym/",
        "whatsapp": "https://wa.me/6281216115104",
        "maps": "https://www.google.com/maps/search/?api=1&query=New+Icon+Gym+Barata+Jaya+59+Surabaya",
        "logo_source": "https://www.instagram.com/new.icon.gym/",
        "facts": [
            "LatihanFisik and Instagram are source basis.",
            "Source notes mention Ruko Plaza Barata Jaya, Jl. Barata Jaya 59 Blok A14, Surabaya, and WhatsApp contact.",
            "Website should emphasize gym floor, aerobic, trainer, pricing inquiry, and Maps.",
        ],
        "hero": {
            "headline": "New Icon Gym",
            "subhead": "A neighborhood Surabaya gym built for practical training, aerobic energy, trainer support, and easy Barata Jaya access.",
            "cta_primary": "Ask New Icon",
            "cta_secondary": "View Barata Jaya",
        },
        "facility_focus": ["gym floor", "aerobic", "trainer support", "local access"],
        "program_focus": ["membership inquiry", "classes", "trainer consultation", "price/pricelist request"],
    },
    {
        "no": 10,
        "name": "SpeedRocky Gym",
        "slug": "speedrocky-gym",
        "city": "Surabaya",
        "archetype": "hardcore",
        "positioning": "Hardcore bodybuilding and strength-focused local gym in Keputih with old-school performance energy.",
        "primary_source": "https://latihanfisik.com/gym-detail/speedrocky-gym/",
        "instagram": "https://www.instagram.com/speedrocky_gym/",
        "whatsapp": "tel:+62315932336",
        "maps": "https://www.google.com/maps/search/?api=1&query=SpeedRocky+Gym+Keputih+Utara+83+Surabaya",
        "logo_source": "https://www.instagram.com/speedrocky_gym/",
        "facts": [
            "LatihanFisik and Instagram location are source basis.",
            "Source notes mention Jl. Keputih Utara No.83, Surabaya and phone 031-5932336.",
            "Avoid wellness styling; use gritty strength visual language.",
        ],
        "hero": {
            "headline": "SpeedRocky Gym",
            "subhead": "An old-school Surabaya strength gym for lifters who want focused equipment, direct access, and serious training energy.",
            "cta_primary": "Call SpeedRocky",
            "cta_secondary": "Open Keputih Maps",
        },
        "facility_focus": ["free weights", "bodybuilding equipment", "strength area", "local training floor"],
        "program_focus": ["bodybuilding", "strength routines", "membership inquiry", "location visit"],
    },
    {
        "no": 11,
        "name": "Warriors Gym Surabaya",
        "slug": "warriors-gym-surabaya",
        "city": "Surabaya",
        "archetype": "hardcore",
        "positioning": "Semi-hardcore strength community gym with warrior-style discipline, private training, and functional training.",
        "primary_source": "https://latihanfisik.com/gym-detail/warriors-gym/",
        "instagram": "https://www.instagram.com/warriorsgym_sby/",
        "whatsapp": "Instagram/DM",
        "maps": "https://www.google.com/maps/search/?api=1&query=Warriors+Gym+Manyar+Adi+II+No+35+Surabaya",
        "logo_source": "https://www.instagram.com/warriorsgym_sby/",
        "facts": [
            "LatihanFisik and Instagram are source basis.",
            "Source notes mention fitness center, private training, functional training, operating hours, and Manyar Adi II No.35.",
            "CTA should route to Instagram/DM until direct WhatsApp is confirmed.",
        ],
        "hero": {
            "headline": "Warriors Gym Surabaya",
            "subhead": "A Surabaya strength community for disciplined training, functional work, and private coaching support.",
            "cta_primary": "Message Warriors",
            "cta_secondary": "Open Manyar Maps",
        },
        "facility_focus": ["strength floor", "functional training", "private training", "community"],
        "program_focus": ["warrior program cards", "PT inquiry", "functional training", "membership"],
    },
    {
        "no": 12,
        "name": "Audid Gym",
        "slug": "audid-gym",
        "city": "Surabaya",
        "archetype": "sport-complex",
        "positioning": "Sport complex in Klampis with gym, court sports, calisthenics, boxing, pool, ice bath, and multi-activity appeal.",
        "primary_source": "https://latihanfisik.com/gym-detail/audid-gym/",
        "instagram": "https://www.instagram.com/audidgym/",
        "whatsapp": "https://wa.me/6282229304618",
        "maps": "https://www.google.com/maps/search/?api=1&query=Audid+Gym+Klampis+Jaya+Surabaya",
        "logo_source": "https://www.instagram.com/audidgym/",
        "facts": [
            "LatihanFisik and Instagram are source basis.",
            "Source notes mention gym, basketball court, pickleball, jogging track, calisthenics, boxing, table tennis, pool, and ice bath.",
            "Primary WhatsApp is 0822-2930-4618; secondary telephone is (031) 5916163.",
        ],
        "hero": {
            "headline": "Audid Gym",
            "subhead": "A complete Surabaya sport complex for gym training, court sports, boxing, recovery, and active community movement.",
            "cta_primary": "Ask Audid",
            "cta_secondary": "View Klampis Maps",
        },
        "facility_focus": ["gym", "basketball court", "pickleball", "jogging track", "calisthenics", "boxing", "table tennis", "pool", "ice bath"],
        "program_focus": ["activity booking", "membership", "training zones", "sport complex tour"],
    },
    {
        "no": 13,
        "name": "Champion Gym Surabaya",
        "slug": "champion-gym-surabaya",
        "city": "Surabaya",
        "archetype": "local-community",
        "positioning": "Budget-comfort gym with full AC, free parking, PT, comfortable training, promo membership, and direct WhatsApp.",
        "primary_source": "https://latihanfisik.com/gym-detail/champion-gym-surabaya/",
        "instagram": "https://www.instagram.com/champion.gym.sby/",
        "whatsapp": "https://wa.me/6281331717779",
        "maps": "https://www.google.com/maps/search/?api=1&query=Champion+Gym+Klampis+Semolo+Timur+6C+Surabaya",
        "logo_source": "https://www.instagram.com/champion.gym.sby/",
        "facts": [
            "LatihanFisik, Instagram, and Waze/location references are source basis.",
            "Source notes mention Klampis Semolo Timur 6C, full AC, free parking, comfortable training, PT, and contact.",
            "Hero should sell comfort and practical value, not fake championship claims.",
        ],
        "hero": {
            "headline": "Champion Gym Surabaya",
            "subhead": "A comfortable Surabaya gym with full AC, free parking, trainer support, and membership promos for everyday training.",
            "cta_primary": "Chat Champion",
            "cta_secondary": "Open Location",
        },
        "facility_focus": ["full AC", "free parking", "comfortable training floor", "personal trainer"],
        "program_focus": ["promo membership", "PT inquiry", "daily/monthly access", "location visit"],
    },
    {
        "no": 14,
        "name": "Crystal Gym & Aerobic",
        "slug": "crystal-gym-aerobic",
        "city": "Surabaya",
        "archetype": "women-wellness",
        "positioning": "Friendly gym and aerobic center in Prapen Surabaya with clean, active, community-class appeal.",
        "primary_source": "Instagram official",
        "instagram": "https://www.instagram.com/newcrystalgym/",
        "whatsapp": "https://wa.me/6282233225222",
        "maps": "https://www.google.com/maps/search/?api=1&query=Crystal+Gym+And+Aerobic+Soponyono+19+Surabaya",
        "logo_source": "https://www.instagram.com/newcrystalgym/",
        "facts": [
            "Official Instagram is source basis; Trip/location references mention fitness and body weight.",
            "Source notes mention Jl. Soponyono No.19 Prapen Surabaya and WhatsApp.",
            "Design should focus on aerobic/class community, not only strength training.",
        ],
        "hero": {
            "headline": "Crystal Gym & Aerobic",
            "subhead": "A friendly Surabaya fitness and aerobic space for movement, body-weight training, classes, and community energy.",
            "cta_primary": "Book Class Info",
            "cta_secondary": "Open Prapen Maps",
        },
        "facility_focus": ["fitness area", "body-weight training", "aerobic studio", "community classes"],
        "program_focus": ["aerobic schedule", "class pass", "membership", "WhatsApp booking"],
    },
    {
        "no": 15,
        "name": "M Gym Malang",
        "slug": "m-gym-malang",
        "city": "Malang",
        "archetype": "local-community",
        "positioning": "Student-friendly local gym near STIE Malangkucecwara/ABM with affordable, simple, comfortable training.",
        "primary_source": "https://latihanfisik.com/gym-detail/m-gym/",
        "instagram": "https://www.instagram.com/mgym.abm/",
        "whatsapp": "Instagram/DM",
        "maps": "https://www.google.com/maps/search/?api=1&query=M+Gym+STIE+Malangkucecwara+ABM+Malang",
        "logo_source": "https://www.instagram.com/mgym.abm/",
        "facts": [
            "LatihanFisik and Instagram are source basis.",
            "Source notes mention Jl. Terusan Candi Kalasan, STIE Malangkucecwara/ABM area, and personal trainer.",
            "Student-friendly language can be used because source notes support campus-area positioning; avoid exact pricing unless verified.",
        ],
        "hero": {
            "headline": "M Gym Malang",
            "subhead": "A practical Malang gym near the ABM campus area for everyday training, PT support, and simple community fitness.",
            "cta_primary": "Message M Gym",
            "cta_secondary": "Open ABM Area Maps",
        },
        "facility_focus": ["daily training", "student area access", "personal trainer", "comfortable local gym"],
        "program_focus": ["membership inquiry", "PT support", "student-friendly visit", "location guidance"],
    },
    {
        "no": 16,
        "name": "DM Gym / DM Fitness Yogyakarta",
        "slug": "dm-gym-yogyakarta",
        "city": "Yogyakarta",
        "archetype": "local-community",
        "positioning": "Jogja community gym with multiple studios, classes, PT, all-branch membership, basketball, and women-friendly area.",
        "primary_source": "https://latihanfisik.com/gym-detail/dm-fitness/",
        "instagram": "https://www.instagram.com/dm_fitnessjogja/ and https://www.instagram.com/dm_gymjogja/",
        "whatsapp": "https://wa.me/6287832893128",
        "maps": "https://www.google.com/maps/search/?api=1&query=DM+Fitness+Jl+Veteran+No+23+Yogyakarta",
        "logo_source": "https://www.instagram.com/dm_fitnessjogja/",
        "facts": [
            "LatihanFisik and Instagram accounts are source basis.",
            "Source notes mention Studio 1-4, WA, women-specific area, all-branch membership, PT, classes, and basketball court.",
            "Website should make branch/studio navigation clear.",
        ],
        "hero": {
            "headline": "DM Fitness Yogyakarta",
            "subhead": "A Jogja training community with studio options, classes, PT, women-friendly area, and all-branch membership direction.",
            "cta_primary": "Chat DM Fitness",
            "cta_secondary": "View Jl. Veteran",
        },
        "facility_focus": ["Studio 1-4", "women-specific area", "basketball court", "personal trainer"],
        "program_focus": ["classes", "all-branch membership", "PT", "branch/studio finder"],
    },
    {
        "no": 17,
        "name": "OCIGEN Fitness",
        "slug": "ocigen-fitness",
        "city": "Yogyakarta",
        "archetype": "premium-wellness",
        "positioning": "Premium hotel fitness center at KJ Hotel with pool, PT, modern equipment, Muay Thai, Pound Fit, Yoga, and Zumba.",
        "primary_source": "https://latihanfisik.com/gym-detail/ocigen-physical-fitness-center/",
        "instagram": "https://www.instagram.com/ocigen.id/",
        "whatsapp": "Instagram/DM",
        "maps": "https://www.google.com/maps/search/?api=1&query=OCIGEN+Physical+Fitness+Center+KJ+Hotel+Yogyakarta",
        "logo_source": "https://www.instagram.com/ocigen.id/",
        "facts": [
            "LatihanFisik is strong source for address, hours, and programs.",
            "Source notes mention KJ Hotel Lt.2, Jl. Parangtritis No.120, hours 06.00-21.00, PT, pool, Muay Thai, Pound Fit, Yoga, and Zumba.",
            "CTA can route to Instagram/DM until verified WhatsApp exists.",
        ],
        "hero": {
            "headline": "OCIGEN Fitness",
            "subhead": "A hotel-based Yogyakarta fitness center with modern training, pool access, group classes, and personal trainer support.",
            "cta_primary": "Message OCIGEN",
            "cta_secondary": "Open KJ Hotel Maps",
        },
        "facility_focus": ["KJ Hotel Lt.2", "pool", "modern equipment", "personal trainer"],
        "program_focus": ["Muay Thai", "Pound Fit", "Yoga", "Zumba", "PT consultation"],
    },
    {
        "no": 18,
        "name": "Optimum Fitness & Cafe",
        "slug": "optimum-fitness-cafe",
        "city": "Yogyakarta",
        "archetype": "cafe-lifestyle",
        "positioning": "Fitness and healthy cafe lifestyle concept in Jogja with training, diet food/cafe, classes, InBody, and membership.",
        "primary_source": "https://latihanfisik.com/gym-detail/optimum-fitness-cafe/",
        "instagram": "https://www.instagram.com/optimumfcyogyakarta/",
        "whatsapp": "Instagram/DM",
        "maps": "https://www.google.com/maps/search/?api=1&query=Optimum+Fitness+Cafe+RW+Monginsidi+39+Yogyakarta",
        "logo_source": "https://www.instagram.com/optimumfcyogyakarta/",
        "facts": [
            "LatihanFisik and Instagram are source basis.",
            "Source notes mention Jl. RW Monginsidi 39, diet food/cafe, classes, InBody test, and membership.",
            "Do not treat it as a normal gym only; cafe/recovery/social space must be a first-class section.",
        ],
        "hero": {
            "headline": "Optimum Fitness & Cafe",
            "subhead": "Train harder, recover better, and enjoy a comfortable cafe space in one active Yogyakarta destination.",
            "cta_primary": "Ask Optimum",
            "cta_secondary": "View Location",
        },
        "facility_focus": ["gym", "diet food/cafe", "InBody test", "class space", "membership"],
        "program_focus": ["fitness packages", "cafe menu/healthy food", "classes", "membership inquiry"],
    },
    {
        "no": 19,
        "name": "BlackBox Gym Bausasran",
        "slug": "blackbox-gym-bausasran",
        "city": "Yogyakarta",
        "archetype": "hardcore",
        "positioning": "Urban dark gym in Bausasran Jogja with 24-hour access, Pro package, Pilates studio, HYROX, sauna, locker, shower, full AC, and PT.",
        "primary_source": "https://www.blackbox.camp/ and https://latihanfisik.com/gym-detail/blackbox-gym-bausasran/",
        "instagram": "https://www.instagram.com/blackbox.camp/",
        "whatsapp": "Instagram/DM",
        "maps": "https://www.google.com/maps/search/?api=1&query=BlackBox+Gym+Bausasran+No+22+Yogyakarta",
        "logo_source": "https://www.instagram.com/blackbox.camp/",
        "facts": [
            "Official website, Instagram, and LatihanFisik are source basis.",
            "Source notes mention Jl. Bausasran No.22, 24-hour access, Pro package, Pilates studio, HYROX class, sauna, locker, shower, full AC, and PT.",
            "Design should feel urban, dark, and serious, but still clean enough for commercial conversion.",
        ],
        "hero": {
            "headline": "BlackBox Gym Bausasran",
            "subhead": "A 24-hour urban training space in Jogja with performance classes, recovery facilities, and serious coaching support.",
            "cta_primary": "Message BlackBox",
            "cta_secondary": "Open Bausasran Maps",
        },
        "facility_focus": ["24-hour gym", "Pilates studio", "HYROX class", "sauna", "locker", "shower", "full AC", "PT"],
        "program_focus": ["Pro package", "performance classes", "Pilates", "PT", "membership"],
    },
    {
        "no": 20,
        "name": "Glanzfit Yogyakarta",
        "slug": "glanzfit-yogyakarta",
        "city": "Yogyakarta",
        "archetype": "women-wellness",
        "positioning": "Women-focused boutique gym in Yogyakarta with private, feminine, clean, and class-based fitness positioning.",
        "primary_source": "Instagram official",
        "instagram": "https://www.instagram.com/glanzfit.id/ and https://www.instagram.com/glanzfit/",
        "whatsapp": "https://wa.me/6287846351704",
        "maps": "https://www.google.com/maps/search/?api=1&query=Glanzfit+Women+Gym+Yogyakarta",
        "logo_source": "https://www.instagram.com/glanzfit.id/",
        "facts": [
            "Official Instagram accounts are source basis.",
            "Source notes mention women's boutique gym, operating hours, WhatsApp, and old/new account references.",
            "Website must emphasize privacy, comfort, class support, and modern women-focused fitness without overclaiming.",
        ],
        "hero": {
            "headline": "Glanzfit Yogyakarta",
            "subhead": "A women-focused boutique gym in Yogyakarta for confidence, privacy, strength, and consistent class-based training.",
            "cta_primary": "Chat Glanzfit",
            "cta_secondary": "Open Maps",
        },
        "facility_focus": ["women-focused gym", "private comfort", "class-based fitness", "modern training space"],
        "program_focus": ["membership inquiry", "class schedule", "consultation", "women-friendly training"],
    },
]


def slugify(text):
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text).strip("-")
    return text


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=TABLE_BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def paragraph_border_bottom(paragraph, color="808080", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=16, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(run, size=13, color=BLUE, bold=True)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(run, size=12, color=BLUE_DARK, bold=True)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    set_table_borders(table, color="E0E6ED")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, 120, 120, 160, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=11, color=BLUE_DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=BLACK)
    doc.add_paragraph()


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table)
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        left, right = table.rows[i].cells
        set_cell_width(left, 1800)
        set_cell_width(right, 7560)
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
        set_cell_shading(left, LIGHT_GRAY)
        p = left.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        set_run_font(r, size=10, color=BLUE_DARK, bold=True)
        p = right.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(value)
        set_run_font(r, size=10)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table)
    set_table_borders(table)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(headers[idx])
        set_run_font(r, size=9.5, color=BLUE_DARK, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_width(cells[idx], widths[idx])
            set_cell_margins(cells[idx])
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=9.5, color=BLACK)
    doc.add_paragraph()
    return table


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")

    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"].font.size = Pt(11)
    styles["List Number"].font.name = "Calibri"
    styles["List Number"].font.size = Pt(11)
    return doc


def add_header_footer(doc, gym):
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    r = header_p.add_run(f"{gym['name']} PRD | {PRESET_NAME}")
    set_run_font(r, size=9, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    r = footer_p.add_run("Commercial landing page product requirements")
    set_run_font(r, size=9, color=MUTED)


def add_title_block(doc, gym):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(r, size=23, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(f"{gym['name']} - Premium Gym Landing Page")
    set_run_font(r, size=14, color=MUTED)

    add_metadata_table(
        doc,
        [
            ("Gym", gym["name"]),
            ("Market", gym["city"]),
            ("Positioning", gym["positioning"]),
            ("Document Version", "v1.0"),
            ("Date", DOC_DATE),
            ("Design Preset", f"{PRESET_NAME} / {HEADER_PATTERN}"),
            ("Status", "Ready for design, implementation, QA, and presentation"),
        ],
    )
    rule = doc.add_paragraph()
    paragraph_border_bottom(rule, color="808080", size="8")


def source_rows(gym):
    return [
        ("Primary source", gym["primary_source"], "Use for verified business facts, facility claims, membership direction, and address context."),
        ("Instagram", gym["instagram"], "Use for logo reference, visual tone, community proof, class/promo cues, and social CTA."),
        ("WhatsApp / Contact", gym["whatsapp"], "Use for primary conversion CTA. If this is Instagram/DM or contact page, do not invent a phone number."),
        ("Maps", gym["maps"], "Use for location CTA and route intent. Embed only if a verified map URL is available."),
        ("Logo source", gym["logo_source"], "Use only to create a local logo asset or clean text fallback. Never hotlink a fragile image."),
    ]


def requirements_by_archetype(gym):
    archetype = ARCHETYPES[gym["archetype"]]
    sections = []
    sections.append((
        "Hero",
        "Instantly communicate the gym name, location, and why this visitor should keep reading.",
        f"Headline: {gym['hero']['headline']}. Subtitle: {gym['hero']['subhead']}. Primary CTA: {gym['hero']['cta_primary']}. Secondary CTA: {gym['hero']['cta_secondary']}.",
        "Hero fits desktop and mobile viewport, uses verified logo/visual direction, and keeps CTA above the fold.",
    ))
    if gym["archetype"] == "chain-modern":
        sections += [
            ("Network / access proposition", "Show scale, convenience, access model, or club ecosystem.", "Use only source-backed claims such as 24/7 access, number of clubs, class count, or facility scale.", "Visitor understands how access/membership works before pricing."),
            ("Facility bento grid", "Make facility depth scannable and visual.", ", ".join(gym["facility_focus"]), "Every card uses real or approved imagery and has concise explanatory copy."),
            ("Program universe", "Present classes, Pilates, PT, recovery, or lifestyle services as selectable tabs.", ", ".join(gym["program_focus"]), "Tabs/buttons are clickable and route to relevant CTA or content."),
            ("Membership comparison", "Convert consideration into action.", "Starter, monthly, club/network, class, PT, or free trial options with placeholders where price is unverified.", "No fake pricing; unknown values are marked 'confirm with admin'."),
        ]
    elif gym["archetype"] == "hardcore":
        sections += [
            ("Training culture", "Differentiate the gym through discipline, strength, and performance.", "Short, direct copy focused on training seriousness and local proof.", "Avoid fake trophy/member claims; tone stays grounded."),
            ("Equipment proof", "Show the training floor and heavy-use areas.", ", ".join(gym["facility_focus"]), "Gallery uses gritty but readable imagery with captions on hover/click."),
            ("Programs / PT", "Make visitors ask about coaching or membership.", ", ".join(gym["program_focus"]), "CTA generates message for strength, PT, membership, or first visit."),
            ("Challenge / transformation proof", "Create motivation without fake before-after claims.", "Use community/gallery proof or 'start your first week' challenge if verified proof is missing.", "Proof labels are honest and do not fabricate outcomes."),
        ]
    elif gym["archetype"] == "women-wellness":
        sections += [
            ("Comfort and confidence", "Explain why the space feels safe, supportive, or class-friendly.", "Use source-backed comfort, class, pool, women-focused, or boutique claims.", "No patronizing copy; design feels premium and clean."),
            ("Facilities / studio", "Show the spaces that reduce friction to visit.", ", ".join(gym["facility_focus"]), "Cards use bright imagery, clear labels, and accessible contrast."),
            ("Classes / instructors", "Make class-based discovery easy.", ", ".join(gym["program_focus"]), "Filter by class type/day if schedule is available; otherwise use request CTA."),
            ("Consultation / membership", "Turn interest into a soft conversion.", "CTA to WhatsApp/DM/Linktree for class info, membership, or consultation.", "Mobile CTA is visible and easy to tap."),
        ]
    elif gym["archetype"] == "cafe-lifestyle":
        sections += [
            ("Train plus recover concept", "Make the hybrid fitness/cafe value obvious.", "Explain gym, cafe/diet food, InBody, classes, and membership based on source notes.", "Visitor understands this is not a normal gym-only page."),
            ("Gym packages", "Present fitness offer first.", ", ".join(gym["program_focus"]), "Pricing can be placeholder if unverified, but CTA must work."),
            ("Cafe / recovery", "Show lifestyle and social reason to visit.", ", ".join(gym["facility_focus"]), "Use warm imagery and descriptive copy, not generic luxury wording."),
            ("Location and reservation", "Connect training, cafe, and maps into one action path.", "Maps + Instagram/DM CTA.", "Visitor can ask admin or open location in one tap."),
        ]
    elif gym["archetype"] == "sport-complex":
        sections += [
            ("Activity ecosystem", "Show the breadth of sports and facilities.", ", ".join(gym["facility_focus"]), "Use category cards or tabs for each sport/activity."),
            ("Facility zone tabs", "Make the complex easier to navigate.", "Gym, court sports, boxing/calisthenics, pool/recovery, track.", "Tabs are keyboard-accessible and mobile-friendly."),
            ("Membership / booking", "Help visitors decide whether to ask about gym, sport, pool, or recovery access.", ", ".join(gym["program_focus"]), "CTA message includes selected activity."),
            ("Gallery and proof", "Show this is a real multi-activity venue.", "Photos and captions for multiple zones.", "No broken images; captions describe each zone."),
        ]
    elif gym["archetype"] == "premium-wellness":
        sections += [
            ("Hotel/wellness value", "Position the space as polished and health-focused.", "Use hotel, pool, class, PT, and modern equipment facts.", "Tone is premium but not fake-luxury."),
            ("Facility and pool proof", "Show why the location feels complete.", ", ".join(gym["facility_focus"]), "Imagery feels calm, clean, and real."),
            ("Classes and PT", "Make class/PT decisions easy.", ", ".join(gym["program_focus"]), "Schedule or request CTA is easy to find."),
            ("Location and hours", "Anchor trust with verified hotel/floor/address/hours.", "KJ Hotel context and hours if applicable.", "Maps CTA works and no unverified phone is invented."),
        ]
    sections += [
        ("Gallery / social proof", "Make the page feel real and alive.", "At least 8 visual assets across hero, facility, class, trainer/community, and location context.", "Every image has alt text, lazy loading where appropriate, and no broken icon."),
        ("Location", "Remove friction for first visit.", "Address, Maps CTA, contact, operating hours if verified.", "Maps button works; unknown details are marked pending verification."),
        ("FAQ", "Answer the last objections before contact.", "Pricing verification, trial/visit, PT, classes, parking, what to bring, schedule changes.", "Accordion works with keyboard and does not hide critical CTA."),
        ("Final CTA", "End with a clear next action.", f"Repeat {gym['hero']['cta_primary']} and {gym['hero']['cta_secondary']} with WhatsApp/Maps links.", "CTA uses verified links and generates a useful message."),
    ]
    return sections


def functional_requirements(gym):
    return [
        ("FR-01", "Mobile navigation", "Open/close menu, trap accidental overflow, close on link click, and show active section.", "No horizontal scroll; all links are reachable by keyboard and tap."),
        ("FR-02", "WhatsApp/contact CTA", f"Primary CTA uses {gym['whatsapp']} and includes selected interest in the message.", "CTA never points to a dead href. Unknown WA uses IG/DM/contact source."),
        ("FR-03", "Maps CTA", f"Secondary CTA opens {gym['maps']}.", "Works on desktop and mobile; label clearly says location/maps."),
        ("FR-04", "Gallery interaction", "Images open in a lightbox or expanded view with short descriptive paragraph visible on hover/click.", "No text blocks permanently cover the photo; Escape closes modal."),
        ("FR-05", "Facility/program tabs", "User can switch between facility categories or program cards based on the gym type.", "Button states are visible: default, hover, active, focus."),
        ("FR-06", "FAQ accordion", "Accordion answers common joining questions without overloading page length.", "Keyboard works and screen readers receive expanded/collapsed state."),
        ("FR-07", "Sticky conversion path", "Use floating WhatsApp/contact CTA on mobile and a clear nav CTA on desktop.", "CTA does not cover content and respects safe area spacing."),
        ("FR-08", "Open/closed status", "If verified hours exist, show current status; otherwise show 'hours pending verification'.", "No fake hours are invented."),
    ]


def content_requirements(gym):
    return [
        ("Hero copy", gym["hero"]["headline"], gym["hero"]["subhead"]),
        ("Primary CTA", gym["hero"]["cta_primary"], f"Route to {gym['whatsapp']} with a prefilled inquiry."),
        ("Secondary CTA", gym["hero"]["cta_secondary"], f"Route to {gym['maps']}."),
        ("Facility copy", "Facility cards must explain the real purpose of each area.", ", ".join(gym["facility_focus"])),
        ("Program copy", "Program cards must explain who each program is for.", ", ".join(gym["program_focus"])),
        ("Data integrity", "Unverified prices, ratings, exact schedules, and exact membership terms must be placeholders.", "Use 'confirm with admin' or 'pending verification' instead of invented claims."),
    ]


def add_source_inventory(doc, gym):
    add_heading(doc, "1. Source Inventory And Data Rules", 1)
    add_paragraph(
        doc,
        "This PRD uses the source table supplied by the project owner as the authoritative brief. The website must prefer official sources first, then Instagram, Linktree/contact pages, Maps, and LatihanFisik or directory sources for local validation.",
    )
    rows = source_rows(gym)
    add_table(doc, ["Source Type", "URL / Reference", "Required Usage"], rows, widths=[1700, 3400, 4260])
    add_callout(
        doc,
        "Data Integrity Rule",
        "Do not invent membership prices, ratings, number of members, trainer certifications, exact schedules, operating hours, or facility claims unless the detail is present in the approved source set. Use a visible 'pending verification' label where needed.",
    )


def add_executive_summary(doc, gym):
    archetype = ARCHETYPES[gym["archetype"]]
    add_heading(doc, "2. Executive Summary", 1)
    add_paragraph(
        doc,
        f"{gym['name']} needs a commercial landing page that feels like a real {archetype['label'].lower()}, not a repeated parallax template. The page should make the visitor understand what the gym is, why it is relevant, where it is located, and what action to take next.",
    )
    add_paragraph(
        doc,
        f"The product objective is to convert qualified visitors into contact actions through {gym['hero']['cta_primary']}, location views through {gym['hero']['cta_secondary']}, and trust-building sections that use verified facts from the approved sources.",
    )
    add_heading(doc, "Key Source-Backed Facts", 2)
    for fact in gym["facts"]:
        add_bullet(doc, fact)


def add_product_goals(doc, gym):
    add_heading(doc, "3. Product Goals And Success Metrics", 1)
    goals = [
        ("G-01", "Increase contact intent", "Visitor can contact the gym from hero, nav, sticky CTA, package/program cards, and final CTA.", "WhatsApp/contact clicks; CTA click-through rate."),
        ("G-02", "Reduce first-visit uncertainty", "Visitor can see location, facilities, programs, and what to ask admin.", "Maps clicks; FAQ engagement; time on page."),
        ("G-03", "Differentiate brand identity", f"Design follows {ARCHETYPES[gym['archetype']]['label']} direction and does not copy the same layout rhythm as other gym pages.", "Qualitative design review; repeated-template audit."),
        ("G-04", "Protect source accuracy", "Only source-backed facts are used; unverified data is clearly marked.", "Content QA pass; zero fake claims."),
        ("G-05", "Prepare production deployment", "Page passes local visual, responsive, link, accessibility, and image checks before deployment.", "Build/validation pass; no broken links/images."),
    ]
    add_table(doc, ["ID", "Goal", "Requirement", "Success Metric"], goals, widths=[900, 1900, 4200, 2360])


def add_users(doc):
    add_heading(doc, "4. Target Users", 1)
    rows = [
        ("First-time visitor", "Wants to know if the gym is close, credible, and beginner-friendly.", "Location, facilities, class/program options, price inquiry, WhatsApp CTA."),
        ("Ready-to-join lead", "Already interested and wants admin response fast.", "Persistent contact CTA, prefilled inquiry, membership/program selector."),
        ("Class or program seeker", "Compares classes, PT, or training style.", "Class/program cards, schedule status, instructor/trainer section if verified."),
        ("Location-driven user", "Searches from Maps or Instagram and needs route clarity.", "Maps CTA, address, hours if verified, parking/access details if source-backed."),
        ("Brand evaluator", "Judges whether the gym looks professional enough to visit.", "Real imagery, logo, polished layout, testimonials/proof where available."),
    ]
    add_table(doc, ["Persona", "Need", "Website Must Answer"], rows, widths=[2100, 3400, 3860])


def add_brand_strategy(doc, gym):
    archetype = ARCHETYPES[gym["archetype"]]
    add_heading(doc, "5. Brand Strategy And Visual Direction", 1)
    rows = [
        ("Positioning", gym["positioning"]),
        ("Brand archetype", archetype["label"]),
        ("Tone", archetype["tone"]),
        ("Visual direction", archetype["visual"]),
        ("Color palette", archetype["palette"]),
        ("Typography", archetype["typography"]),
        ("Image direction", "Use real gym, facility, class, equipment, trainer/community, and location context. Avoid random stock images that conflict with the gym type."),
        ("Avoid", "Repeated template blocks, fake luxury wording, over-neon gaming effects, broken logos, unverified claims, and hover-only critical information."),
    ]
    add_table(doc, ["Brand Area", "Requirement"], rows, widths=[2200, 7160])


def add_information_architecture(doc, gym):
    archetype = ARCHETYPES[gym["archetype"]]
    add_heading(doc, "6. Recommended Page Structure", 1)
    add_paragraph(
        doc,
        "The page structure should follow the business journey below. Keep the reusable component system, but vary section order, emphasis, color, copy, and interaction per gym so the output does not feel factory-made.",
    )
    rows = []
    for idx, section in enumerate(archetype["journey"], start=1):
        rows.append((str(idx), section, "Use gym-specific content and source-backed facts."))
    add_table(doc, ["Order", "Section", "Implementation Note"], rows, widths=[800, 4300, 4260])


def add_section_requirements(doc, gym):
    add_heading(doc, "7. Section-Level Requirements", 1)
    add_table(
        doc,
        ["Section", "Purpose", "Required Content", "Acceptance Criteria"],
        requirements_by_archetype(gym),
        widths=[1700, 2500, 3260, 1900],
    )


def add_functional_requirements(doc, gym):
    add_heading(doc, "8. Functional Requirements", 1)
    add_table(doc, ["ID", "Feature", "Requirement", "Acceptance Criteria"], functional_requirements(gym), widths=[900, 1900, 4260, 2300])


def add_content_and_copy(doc, gym):
    add_heading(doc, "9. Content And Copywriting Requirements", 1)
    add_table(doc, ["Content Area", "Approved Copy / Direction", "Notes"], content_requirements(gym), widths=[1800, 3800, 3760])
    add_heading(doc, "WhatsApp / Contact Message Template", 2)
    message = (
        f"Halo admin {gym['name']}, saya ingin bertanya tentang: Paket/Program: [pilihan user], "
        "Tujuan latihan: [tujuan user], Jadwal kunjungan: [hari/jam pilihan], Kebutuhan: [membership/PT/classes/trial]. "
        "Mohon informasinya. Terima kasih."
    )
    add_callout(doc, "Generated Message Requirement", message)


def add_assets(doc, gym):
    add_heading(doc, "10. Asset And Image Requirements", 1)
    add_paragraph(
        doc,
        "All production assets should be stored locally in the website repository. Do not hotlink fragile external image paths. If an official logo cannot be legally or technically downloaded, use a clean text logo fallback without showing a broken image icon.",
    )
    rows = [
        ("Logo", f"/public/gyms/{gym['slug']}/logo.png", f"Source: {gym['logo_source']}. Must have text fallback."),
        ("Hero", f"/public/gyms/{gym['slug']}/hero.webp", "Real facility, training, or brand-appropriate image. Eager load above the fold."),
        ("Gallery 1", f"/public/gyms/{gym['slug']}/facility-1.webp", "Facility/equipment image with descriptive alt text."),
        ("Gallery 2", f"/public/gyms/{gym['slug']}/facility-2.webp", "Class, PT, or community image with hover/click paragraph."),
        ("Gallery 3", f"/public/gyms/{gym['slug']}/facility-3.webp", "Location, recovery, pool, cafe, or sport-zone context."),
        ("Social proof", f"/public/gyms/{gym['slug']}/community-1.webp", "Member/community proof if available; otherwise approved generic local gym image."),
        ("OG image", f"/public/gyms/{gym['slug']}/og-image.webp", "Used for Open Graph preview and social sharing."),
    ]
    add_table(doc, ["Asset", "Target Path", "Requirement"], rows, widths=[1800, 3400, 4160])


def add_seo_accessibility_performance(doc, gym):
    add_heading(doc, "11. SEO, Accessibility, And Performance", 1)
    seo_rows = [
        ("SEO title", f"{gym['name']} | Gym Landing Page", "Unique per gym; include city if useful."),
        ("Meta description", gym["hero"]["subhead"], "Keep concise and source-backed."),
        ("Canonical", f"https://{gym['slug']}.vercel.app/ or final custom domain", "Update when production domain is confirmed."),
        ("Schema", "HealthClub / SportsActivityLocation / LocalBusiness", "Include address, sameAs, hours, and phone only when verified."),
        ("Open Graph", "Title, description, image, URL", "Use local OG image path."),
    ]
    add_table(doc, ["Item", "Requirement", "Notes"], seo_rows, widths=[1800, 4200, 3360])
    add_heading(doc, "Accessibility Requirements", 2)
    for item in [
        "One H1 only; heading order must be logical from H1 to H2/H3.",
        "All important images require descriptive alt text.",
        "CTA buttons need visible focus states, hover states, and minimum mobile tap size.",
        "Modal/lightbox must close with Escape and should not trap users without a close control.",
        "Color contrast must remain readable; do not rely only on color to indicate active state.",
        "Respect prefers-reduced-motion and simplify parallax on mobile.",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "Performance Requirements", 2)
    for item in [
        "Lazy-load below-the-fold images and use explicit aspect ratios to reduce layout shift.",
        "Avoid huge image files; use optimized WebP/AVIF where practical.",
        "Keep JavaScript small and defensive; no console errors if optional sections are missing.",
        "Do not add heavy animation libraries unless already available and justified.",
    ]:
        add_bullet(doc, item)


def add_conversion_and_analytics(doc, gym):
    add_heading(doc, "12. Conversion Flow And Analytics", 1)
    rows = [
        ("Step 1", "Visitor lands from Instagram, search, Maps, or shared link.", "Hero immediately shows gym identity, location, CTA, and visual proof."),
        ("Step 2", "Visitor scans facilities/programs.", "Bento grid, tabs, and image captions answer whether this gym fits them."),
        ("Step 3", "Visitor selects interest.", "Membership/program/facility CTA stores selected interest for WhatsApp/contact message."),
        ("Step 4", "Visitor clicks WhatsApp/contact or Maps.", "Message/route opens correctly with no dead button."),
        ("Step 5", "Admin receives better-qualified lead.", "Message includes package/program, goal, visit time, and need."),
    ]
    add_table(doc, ["Flow", "User Action", "Website Response"], rows, widths=[1200, 3400, 4760])
    add_heading(doc, "Analytics Events", 2)
    for item in [
        "cta_whatsapp_click",
        "cta_maps_click",
        "membership_select",
        "program_select",
        "gallery_open",
        "faq_expand",
        "instagram_click",
    ]:
        add_bullet(doc, item)


def add_acceptance_criteria(doc, gym):
    add_heading(doc, "13. Acceptance Criteria", 1)
    criteria = [
        "Website visually matches the gym's brand archetype and does not look identical to the other gym sites.",
        "Hero logo does not appear broken; if logo asset is unavailable, clean text fallback appears.",
        "Hero title and CTA fit inside mobile and desktop viewport without horizontal overflow.",
        "Primary contact CTA and Maps CTA work from nav, hero, sticky CTA, and final CTA.",
        "Facility/program/gallery sections contain gym-specific copy and images or approved fallback images.",
        "No unverified claims about rating, number of members, exact price, exact schedule, or certifications are published.",
        "FAQ accordion, gallery lightbox, mobile nav, active states, and any selectors work without console errors.",
        "SEO metadata, Open Graph tags, JSON-LD, alt text, and canonical target are present.",
        "Accessibility checks pass for keyboard navigation, focus states, readable contrast, and reduced motion.",
        "Local build/validation passes before GitHub push and Vercel deployment.",
    ]
    for item in criteria:
        add_bullet(doc, item)


def add_risks_and_roadmap(doc, gym):
    add_heading(doc, "14. Risks, Assumptions, And Roadmap", 1)
    risk_rows = [
        ("Missing price or schedule", "Some sources may not expose complete pricing or live class schedules.", "Use 'confirm with admin' and CTA until official data is provided."),
        ("Logo rights and quality", "Instagram logos may be low resolution.", "Use official website logo when available; otherwise clean vector/text recreation only with approval."),
        ("Image availability", "Official facility photos may require manual collection.", "Use approved local image assets and document placeholders clearly."),
        ("Contact uncertainty", "Some gyms only provide IG/DM or contact page.", "Do not invent WhatsApp. Route CTA to verified contact channel."),
    ]
    add_table(doc, ["Risk", "Impact", "Mitigation"], risk_rows, widths=[2200, 3400, 3760])
    add_heading(doc, "Implementation Roadmap", 2)
    roadmap = [
        ("Phase 1", "Confirm sources and assets", "Download/store logo, hero, gallery, and source screenshots where allowed."),
        ("Phase 2", "Design system", "Set palette, typography, spacing, component style, responsive rules, and CTA states."),
        ("Phase 3", "Landing page build", "Implement hero, facilities, membership/program, gallery, location, FAQ, and final CTA."),
        ("Phase 4", "Interaction QA", "Test nav, tabs, accordion, lightbox, WhatsApp generator, Maps links, and reduced motion."),
        ("Phase 5", "Production release", "Run build/validation, push GitHub, deploy Vercel, and verify live page."),
    ]
    add_table(doc, ["Phase", "Workstream", "Output"], roadmap, widths=[1300, 2700, 5360])


def add_appendix(doc, gym):
    add_heading(doc, "15. Appendix: Approved Source Links", 1)
    rows = source_rows(gym)
    table = doc.add_table(rows=1, cols=3)
    set_table_width(table)
    set_table_borders(table)
    headers = ["Source", "Link", "Usage"]
    widths = [1700, 4000, 3660]
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(headers[idx])
        set_run_font(r, size=9.5, color=BLUE_DARK, bold=True)
    for label, url, usage in rows:
        cells = table.add_row().cells
        for idx, width in enumerate(widths):
            set_cell_width(cells[idx], width)
            set_cell_margins(cells[idx])
        cells[0].paragraphs[0].add_run(label)
        set_run_font(cells[0].paragraphs[0].runs[0], size=9.5, bold=True)
        p = cells[1].paragraphs[0]
        if url.startswith("http"):
            add_hyperlink(p, url, url)
        else:
            r = p.add_run(url)
            set_run_font(r, size=9.5)
        r = cells[2].paragraphs[0].add_run(usage)
        set_run_font(r, size=9.5)


def add_cover_page(doc, gym):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    set_run_font(r, size=22, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(20)
    r = p.add_run(f"Gym Growth Platform - {gym['name']} Listing & Landing Page")
    set_run_font(r, size=16, color=BLUE_DARK, bold=True)

    add_metadata_table(
        doc,
        [
            ("Product Name", f"Gym Growth Platform - {gym['name']}"),
            ("Product Scope", "Public gym directory, commercial landing page, gym dashboard, monetization, and admin operations"),
            ("Version", "v1.1"),
            ("Status", "Draft for product, design, engineering, SEO, and business review"),
            ("Date", DOC_DATE),
            ("Primary Reviewer", "Albert Saputra"),
            ("Reviewers", "Product Owner, UI/UX Designer, Frontend Engineer, SEO/CRO Reviewer, Gym Business Reviewer"),
            ("Target Launch Market", gym["city"]),
        ],
    )

    add_callout(
        doc,
        "Document Intent",
        "This PRD defines the product requirements for a scalable gym discovery and lead-generation platform, with this document tailored to the listed gym as one rollout candidate from the 20-gym dataset.",
    )
    doc.add_page_break()


def dataset_key_metrics_rows():
    return [
        ("Monthly Active Users", "15,000 MAU by month 6", "Public directory users browsing gym pages."),
        ("Organic Search Traffic", "10,000 organic sessions/month by month 6", "SEO landing pages for city, gym, facility, and class intent."),
        ("WhatsApp Leads", "1,200 WA/contact clicks/month by month 6", "Tracked from hero, sticky CTA, package cards, and final CTA."),
        ("Gym Claim Rate", "30% of initial 20 gyms claimed by month 4", "Gyms confirm data and activate dashboard access."),
        ("Listing Completeness", "90% verified core data coverage by month 3", "Logo, city, address/maps, contact, facilities, programs, and source notes."),
        ("Paid Conversion", "10 paid Starter/Pro gyms by month 6", "Starter/Pro monetization from claimed listings and premium landing pages."),
    ]


def city_breakdown_rows():
    return [
        ("Malang", "4", "2", "4", "2 direct WA + 1 Linktree", "4", "4", "Strong for local gym + campus + wellness positioning."),
        ("Surabaya", "8", "0 official gym websites in sample", "8", "7 direct WA/tel + 1 IG/DM", "8", "8", "Largest city cluster; best MVP supply for SEO directory pages."),
        ("Yogyakarta", "5", "1 official website in sample", "5", "2 direct WA/tel + 3 IG/DM", "5", "5", "Good for boutique, cafe, hotel, and community gym segmentation."),
        ("National / Chain", "3", "3", "3", "3 contact paths", "3", "3", "Use as benchmark brands for feature depth and monetization standards."),
    ]


def benchmark_rows():
    return [
        ("Calm", "SaaS - Health and wellness", "Emotion-led hero, soft trust, simple subscription CTA.", "Use for wellness, women-friendly, hotel fitness, and recovery-focused gym pages."),
        ("Zola", "E-commerce - Wedding", "Guided onboarding, planning checklist, marketplace + service bundle.", "Use for gym onboarding checklist: choose goal, location, program, contact admin."),
        ("CD Baby", "SaaS - Entertainment", "Creator-first benefits, pricing clarity, distribution workflow.", "Use for gym owner dashboard value: publish profile, get leads, track performance."),
        ("Netflix", "SaaS - Entertainment", "Content discovery rows, personalization, subscription comparison.", "Use for class/program discovery rows and recommended gym lists."),
        ("LinkedIn", "SaaS - Professional services", "Profile completeness, credibility graph, network effects.", "Use for claimed gym profile completeness and owner credibility badges."),
        ("Goby", "E-commerce - Health and fitness", "Product education, subscription convenience, clean comparison.", "Use for membership plan and facility benefit explanation."),
        ("DoorDash", "SaaS - Food delivery", "Local search, merchant pages, fast CTA, location availability.", "Use for city-based gym discovery and Maps-first local intent."),
        ("SEMrush", "SaaS - Marketing", "Tool suite packaging, feature depth, lead capture.", "Use for Pro dashboard modules: analytics, SEO, lead insights."),
        ("Coco Village", "E-commerce - Furniture", "Lifestyle visuals, product categories, warm editorial tone.", "Use for family-friendly or lifestyle-oriented gym galleries."),
        ("Grass Roots", "E-commerce - Food and nutrition", "Ingredient/proof storytelling, subscription trust.", "Use for health cafe, recovery, supplement, and diet-food modules."),
        ("Amazon", "SaaS and E-commerce", "Search, filters, reviews, recommendations, trust at scale.", "Use as north star for directory filtering, comparison, reviews, and conversion paths."),
        ("Branch Furniture", "E-commerce - Furniture", "B2B ecommerce, office setup bundles, practical product education.", "Use for corporate gym packages and facility bundle pages."),
        ("Western Rise", "E-commerce - Apparel", "Performance benefit copy, product specs, lifestyle proof.", "Use for performance gym pages: function, durability, training outcomes."),
        ("Athabasca University", "Education", "Program catalog, eligibility clarity, flexible learning journey.", "Use for class/program catalog and beginner pathway logic."),
        ("Bariatric Eating", "Food and nutrition", "Educational content, transformation support, trust-sensitive copy.", "Use for weight-loss goal pages with safe, non-medical disclaimers."),
        ("blow LTD.", "Beauty", "Booking flow, service menu, location/time selection.", "Use for PT trial, class booking, or consultation request flow."),
        ("Blue Forest Farms", "E-commerce - Cannabis", "Compliance-aware product education and trust signals.", "Use as reminder for sensitive claims: no fake health or transformation promises."),
        ("Border Buddy", "Travel and shipping", "Complex process simplified into clear steps.", "Use for explaining listing claim, verification, and gym onboarding."),
        ("Rover", "Pet services", "Marketplace matching, profile trust, reviews, availability.", "Use for matching users to gyms and trainer profiles."),
        ("Campaign Monitor", "SaaS - Marketing", "Use-case pages, templates, conversion education.", "Use for owner-facing marketing automation and lead follow-up value."),
        ("Class Creator", "SaaS - Education", "Class management, roster, admin workflows.", "Use for gym class schedule management in dashboard."),
        ("Fast Mask", "E-commerce - Apparel", "Urgent purchase clarity, product trust, shipping visibility.", "Use for promo modules without fake scarcity."),
        ("Quality Eggs", "E-commerce - Food and nutrition", "Freshness, source proof, simple product value.", "Use for source-backed facility/photo proof and authenticity messaging."),
        ("HomeLoanGurus", "SaaS - Finance and insurance", "Lead form qualification and trust-heavy decision journey.", "Use for high-intent gym consultation and PT inquiry forms."),
        ("Jet Pet", "Pet services", "Premium service trust, facility tours, care reassurance.", "Use for premium gym facility tour and comfort reassurance."),
        ("Mooala", "E-commerce - Food and nutrition", "Brand personality, product benefits, retail locator.", "Use for local gym personality and branch/location finder."),
        ("NANOR", "E-commerce - Health and gifts", "Giftable wellness product, simple catalogue, visual warmth.", "Use for wellness gift/pass or referral campaign concepts."),
        ("Panda7", "SaaS - Finance and insurance", "Quote comparison, simple decision flow, price transparency.", "Use for membership comparison and quote/pricelist request UX."),
        ("Lyft", "Transportation", "Location-based matching, instant CTA, service availability.", "Use for 'gym near me' route, distance, and Maps-led discovery."),
        ("Perfect Keto", "E-commerce - Food and nutrition", "Education-first ecommerce, recipes, goal-specific journeys.", "Use for goal-based training plans and content SEO."),
        ("Gusto", "SaaS - HR", "Friendly B2B SaaS onboarding, payroll complexity simplified.", "Use for owner dashboard onboarding and package education."),
        ("Roomeze", "SaaS - Real estate", "Matching, roommate/profile trust, location filters.", "Use for matching gyms by vibe, gender comfort, class, location."),
        ("Smalls", "E-commerce - Food and nutrition", "Quiz-first purchase flow and personalized plan.", "Use for gym goal quiz: fat loss, strength, class, women-friendly, PT."),
        ("Sundae", "SaaS - Real estate", "Trust-heavy marketplace, seller education, lead capture.", "Use for gym owner sales pages and Pro plan pitch."),
        ("Wavehuggers", "Travel and recreation", "Experience booking, class/session inventory, social proof.", "Use for class, trial pass, and activity booking patterns."),
        ("Woolx", "E-commerce - Apparel", "Material education, performance proof, comparison.", "Use for facility/equipment education and training benefit proof."),
        ("Zumba", "Health and fitness", "Community movement, instructor network, class discovery.", "Use for aerobic/Zumba/class-heavy gym pages."),
        ("Mailchimp", "SaaS - Marketing", "Freemium tiers, playful education, small business tooling.", "Use for Free/Starter/Pro package clarity and owner dashboard tone."),
        ("Spotify", "Streaming audio", "Personalized discovery, playlists, saved items.", "Use for saved gyms, recommended classes, and personalized city pages."),
        ("Snackpass", "E-commerce - Food and nutrition", "Social ordering, campus/local community loops.", "Use for student-friendly gyms, referrals, and community promos."),
    ]


def benchmark_pattern_rows():
    return [
        ("Marketplace discovery", "Amazon, Rover, DoorDash, Lyft, Roomeze", "Search, filters, profile trust, location-first browsing, recommendation rows."),
        ("SaaS dashboard and owner tooling", "SEMrush, Mailchimp, Gusto, Campaign Monitor, Class Creator", "Owner dashboard, analytics, listing completeness, package tiers, admin workflows."),
        ("Subscription and monetization", "Netflix, Calm, Goby, Perfect Keto, CD Baby", "Free/Starter/Pro tier logic, trial CTA, benefits comparison, recurring value."),
        ("Lifestyle and trust storytelling", "Coco Village, Branch Furniture, Western Rise, Jet Pet, Mooala", "Premium visual rhythm, proof cards, editorial content, facility storytelling."),
        ("Education and guided journeys", "Athabasca University, Bariatric Eating, Border Buddy, Quality Eggs", "Goal-based content, process education, trust-sensitive disclaimers."),
        ("Booking and service flow", "blow LTD., Wavehuggers, Zumba, HomeLoanGurus", "Consultation, class booking, qualified lead flow, schedule-aware CTA."),
        ("Community and local loops", "Snackpass, Spotify, LinkedIn", "Saved pages, profile completeness, social proof, referrals, network effect."),
    ]


def add_prd_executive_summary(doc, gym):
    add_heading(doc, "1. Ringkasan Eksekutif", 1)
    add_paragraph(
        doc,
        "Gym Growth Platform adalah platform direktori, landing page komersial, dashboard SaaS, dan marketplace lead-generation untuk membantu pencari gym menemukan tempat latihan yang relevan, sekaligus membantu pemilik gym mendapatkan leads WhatsApp, traffic lokal, profil brand yang lebih profesional, dan data performa listing.",
    )
    add_paragraph(
        doc,
        f"Untuk {gym['name']}, platform ini berfungsi sebagai listing publik dan landing page siap jual yang menampilkan source-backed data, visual brand, fasilitas, program, Maps, Instagram, dan CTA kontak yang valid.",
    )
    add_table(doc, ["Metric", "Target", "Definition"], dataset_key_metrics_rows(), widths=[2400, 2600, 4360])


def add_prd_background(doc, gym):
    add_heading(doc, "2. Latar Belakang", 1)
    add_paragraph(
        doc,
        "Dataset awal berisi 20 gym di Malang, Surabaya, Yogyakarta, dan brand nasional. Temuan utama dari dataset: hanya sekitar 40% gym memiliki website atau halaman brand resmi yang cukup kuat untuk kebutuhan komersial, sementara 0% tercatat memiliki landing page SEO lokal yang jelas untuk ranking Google pada intent seperti 'gym dekat saya', 'personal trainer', 'kelas aerobic', atau 'gym wanita'.",
    )
    add_paragraph(
        doc,
        "Masalah pasar: gym lokal sering memiliki Instagram aktif dan kontak WhatsApp, tetapi tidak punya struktur web yang rapi untuk SEO, trust, fasilitas, membership, jadwal kelas, dan conversion flow. Akibatnya, calon member harus bertanya manual, pemilik gym kehilangan traffic organik, dan penawaran sulit dibandingkan.",
    )
    add_paragraph(
        doc,
        f"Peluang untuk {gym['name']}: membuat halaman yang memadukan data resmi, CTA langsung, Maps, source Instagram, dan struktur produk yang bisa di-scale ke dashboard serta monetisasi.",
    )
    add_heading(doc, "Breakdown Kelengkapan Data Per Kota", 2)
    add_table(
        doc,
        ["Kota / Cluster", "Gym", "Website Resmi", "Instagram", "Kontak", "Maps", "Logo Source", "Catatan Peluang"],
        city_breakdown_rows(),
        widths=[1400, 700, 1500, 1100, 1500, 700, 1000, 1460],
    )
    add_heading(doc, "Benchmark Lintas Industri Untuk Memperluas Produk", 2)
    add_paragraph(
        doc,
        "Agar produk tidak terlalu sempit sebagai katalog gym saja, platform perlu mengambil pola dari SaaS, e-commerce, marketplace, pendidikan, transportasi, dan layanan lokal. Benchmark berikut bersifat directional: digunakan untuk pola halaman, UX, monetisasi, dan informasi produk, bukan untuk menyalin visual brand.",
    )
    add_table(
        doc,
        ["Pattern Family", "Representative Examples", "Adaptasi Untuk Gym Growth Platform"],
        benchmark_pattern_rows(),
        widths=[2200, 3100, 4060],
    )
    add_heading(doc, "Reference Library Halaman", 2)
    add_table(
        doc,
        ["Brand / Example", "Kategori", "Pattern Yang Dipelajari", "Adaptasi Untuk Platform Gym"],
        benchmark_rows(),
        widths=[1700, 2100, 2700, 2860],
    )


def add_prd_okrs(doc):
    add_heading(doc, "3. Tujuan & OKRs", 1)
    rows = [
        ("Adoption", "Membuat pemilik gym percaya dan mau mengklaim listing.", "KR1: 6 dari 20 gym mengklaim listing dalam 4 bulan. KR2: 12 gym menyetujui data profil terverifikasi. KR3: 80% halaman memiliki logo dan foto lokal/fasilitas."),
        ("Traffic", "Membangun traffic organik lokal dari pencari gym.", "KR1: 10,000 organic sessions/month bulan ke-6. KR2: 60 landing pages SEO terindeks. KR3: CTR search minimal 3%."),
        ("Lead Gen", "Mengubah visitor menjadi WhatsApp/contact lead yang siap ditindaklanjuti.", "KR1: 1,200 WA/contact clicks/month bulan ke-6. KR2: conversion CTA minimal 8%. KR3: 70% lead menyertakan kebutuhan/program melalui generator pesan."),
        ("Monetisasi", "Menguji paket listing berbayar untuk gym.", "KR1: 10 paid gyms bulan ke-6. KR2: 3 paket Pro aktif. KR3: churn paid listing di bawah 10% per bulan setelah pilot."),
    ]
    add_table(doc, ["Objective", "Goal", "Key Results"], rows, widths=[1500, 2800, 5060])


def add_prd_personas(doc):
    add_heading(doc, "4. User Personas & Stakeholders", 1)
    persona_rows = [
        ("Primary need", "Cari gym yang dekat, cocok dengan tujuan, jelas fasilitasnya, dan mudah dihubungi.", "Mendapatkan calon member dari traffic Google/Instagram tanpa harus menjawab pertanyaan dasar berulang-ulang."),
        ("Pain point", "Informasi gym tersebar di IG, Maps, dan chat; harga/jadwal/fasilitas tidak selalu jelas.", "Belum punya website SEO, konten tidak terstruktur, sulit membuktikan fasilitas dan promo."),
        ("Decision trigger", "Foto asli, Maps, WhatsApp cepat, class/program jelas, review/proof, dan harga atau starting price.", "Lead WhatsApp naik, tampilan brand lebih profesional, data bisa diedit sendiri."),
        ("Core action", "Search, compare, open Maps, click WhatsApp, save/share gym.", "Claim listing, update data, upload foto, lihat leads/analytics, upgrade paket."),
        ("Success metric", "Menemukan gym yang cocok dan bisa bertanya dalam kurang dari 2 menit.", "Mendapat leads yang lebih siap beli dan halaman yang terlihat layak jual."),
    ]
    add_table(doc, ["Aspect", "Persona A: Pencari Gym", "Persona B: Pemilik Gym"], persona_rows, widths=[1700, 3830, 3830])
    add_heading(doc, "Stakeholder Map", 2)
    stakeholder_rows = [
        ("Product Owner", "Menentukan prioritas fitur, positioning, dan target monetisasi.", "Weekly product review."),
        ("Gym Owner/Admin", "Memvalidasi data, paket, foto, promo, dan kontak.", "Dashboard, WhatsApp, onboarding call."),
        ("Visitor / Calon Member", "Mencari gym, membandingkan pilihan, dan menghubungi admin.", "Public website, SEO pages, mobile CTA."),
        ("Designer", "Membuat sistem visual premium dan brand-specific.", "Design system, responsive prototype."),
        ("Engineer", "Membangun platform, dashboard, SEO, dan integrasi.", "Repo, issue tracker, deployment."),
        ("Admin Platform", "Moderasi data, verifikasi listing, dan support monetisasi.", "Admin panel and reports."),
    ]
    add_table(doc, ["Stakeholder", "Role", "Touchpoint"], stakeholder_rows, widths=[2100, 4360, 2900])


def add_prd_scope_roadmap(doc):
    add_heading(doc, "5. Scope & Roadmap", 1)
    scope_rows = [
        ("In Scope MVP", "Public gym directory, individual gym landing pages, WhatsApp CTA generator, Maps CTA, Instagram/source links, facility/program cards, FAQ, admin-curated data, basic dashboard, analytics events, paid listing tiers."),
        ("Out of Scope MVP", "Full marketplace booking, in-app payments for memberships, live class seat inventory, AI trainer planning, native mobile app, review moderation at scale, multi-language CMS, automated web scraping without approval."),
    ]
    add_table(doc, ["Scope", "Items"], scope_rows, widths=[2200, 7160])
    roadmap_rows = [
        ("Phase 1", "Discovery & Data Validation", "Weeks 1-2", "20 gym dataset cleaned, source confidence tagged, logo/photo checklist ready."),
        ("Phase 2", "MVP Public Directory", "Weeks 3-6", "Searchable public pages, gym profile template, CTA and Maps flow live."),
        ("Phase 3", "Dashboard Pilot", "Weeks 7-10", "Claim listing, edit data, upload photos, view basic leads."),
        ("Phase 4", "Monetization Pilot", "Weeks 11-14", "Free/Starter/Pro packages, payment/manual invoice process, premium page modules."),
        ("Phase 5", "SEO Scale & Ops", "Weeks 15-20", "City/category SEO pages, admin workflows, analytics reporting, production QA."),
    ]
    add_table(doc, ["Phase", "Name", "Target", "Release Output"], roadmap_rows, widths=[1000, 2300, 1400, 4660])


def add_prd_functional(doc, gym):
    add_heading(doc, "6. Kebutuhan Fungsional", 1)
    module_rows = [
        ("Direktori Publik", "Search, city/category filters, gym cards, gym detail pages, Maps/WA CTA, SEO pages.", "User can find and compare gyms quickly."),
        ("Gym Dashboard", "Claim listing, edit profile, update facilities/classes, upload logo/photos, see leads/analytics.", "Gym owner can maintain data without developer help."),
        ("Monetisasi", "Free/Starter/Pro package gating, featured listing, premium landing page, analytics report, lead tracking.", "Platform can sell clear value without disrupting free discovery."),
        ("Admin Panel", "Moderate gym data, approve claims, verify sources, manage packages, publish pages, inspect reports.", "Operations team can maintain quality and prevent fake claims."),
    ]
    add_table(doc, ["Module", "Core Features", "Outcome"], module_rows, widths=[2100, 4960, 2300])
    add_heading(doc, "3-Tier Monetization Package", 2)
    tier_rows = [
        ("Free", "Basic listing, Maps link, Instagram link, limited facility tags, contact CTA.", "Unclaimed or entry-level gyms.", "IDR 0"),
        ("Starter", "Verified listing, photo gallery, WhatsApp lead tracking, basic analytics, promo/class modules.", "Local gyms that want better leads.", "Monthly paid package, final price TBD."),
        ("Pro", "Premium landing page, featured placement, custom SEO sections, campaign tracking, priority updates, monthly report.", "Gyms ready to use the platform as marketing channel.", "Higher monthly package, final price TBD."),
    ]
    add_table(doc, ["Tier", "Features", "Best For", "Pricing Rule"], tier_rows, widths=[1200, 4360, 2300, 1500])
    add_callout(
        doc,
        f"{gym['name']} MVP Fit",
        f"Recommended starting package for this gym: create a source-backed public listing first, then offer Starter/Pro once the gym validates contact flow and image assets.",
    )


def add_prd_nonfunctional(doc):
    add_heading(doc, "7. Non-Fungsional", 1)
    rows = [
        ("Core Web Vitals", "LCP <= 2.5s, INP <= 200ms, CLS <= 0.1", "Use optimized images, stable dimensions, minimal JS."),
        ("Scalability", "Support 1,000+ gym listings and 100k monthly visits", "Use indexed database, cached pages, and CDN deployment."),
        ("Security", "Protect dashboard/admin routes, validate uploads, prevent spam leads.", "Auth, role-based access, file validation, rate limiting."),
        ("SEO", "Indexable pages, schema markup, canonical URLs, sitemap, city/category pages.", "No blocked crawling; unique metadata per gym."),
        ("Accessibility", "WCAG AA target for contrast, keyboard navigation, alt text, focus states.", "Manual QA plus automated checks."),
        ("Browser Compatibility", "Latest Chrome, Safari, Edge, Firefox, and mobile WebView.", "Progressive enhancement; no desktop-only interactions."),
        ("Reliability", "99.5% uptime target for MVP landing pages.", "Static/edge delivery and simple fallback states."),
        ("Data Integrity", "Only publish verified or clearly labeled placeholder data.", "Source confidence fields and admin moderation."),
    ]
    add_table(doc, ["Area", "Target Metric / Requirement", "Implementation Notes"], rows, widths=[2100, 3300, 3960])


def add_prd_model_data(doc):
    add_heading(doc, "8. Model Data", 1)
    rows = [
        ("gyms", "id, slug, name, city, province, address, maps_url, instagram_url, whatsapp_url, website_url, logo_url, status, source_confidence, package_tier, created_at, updated_at", "Main business profile and routing data."),
        ("gym_facilities", "id, gym_id, facility_key, facility_label, description, icon, verified, sort_order", "Facilities such as cardio, free weights, pool, sauna, cafe, classes, parking."),
        ("gym_photos", "id, gym_id, photo_url, photo_type, alt_text, caption, source_url, verified, sort_order", "Hero, logo, gallery, class, trainer/community, location, OG image."),
        ("gym_analytics", "id, gym_id, event_name, event_value, source_page, visitor_city, device_type, occurred_at", "CTA clicks, Maps clicks, gallery opens, package selections, dashboard reporting."),
    ]
    add_table(doc, ["Table", "Fields", "Purpose"], rows, widths=[1800, 4960, 2600])


def add_prd_user_stories(doc):
    add_heading(doc, "9. User Stories & Acceptance Criteria", 1)
    rows = [
        ("US-01", "As a gym seeker, I want to search by city so I can find gyms near me.", "Given I select a city, when results load, then only relevant gym cards appear with Maps and contact CTA."),
        ("US-02", "As a gym seeker, I want to compare facilities before contacting a gym.", "Given I open a gym card/page, then I can see verified facility tags, images, and descriptions."),
        ("US-03", "As a gym seeker, I want one-tap WhatsApp/contact.", "Given I click CTA, then the message includes gym name, selected interest, and visit intent."),
        ("US-04", "As a gym owner, I want to claim my listing.", "Given I submit a claim, then admin can review source/contact proof before dashboard activation."),
        ("US-05", "As a gym owner, I want to update photos and facilities.", "Given I have dashboard access, then I can submit edits that enter moderation before publish."),
        ("US-06", "As a gym owner, I want to see lead performance.", "Given leads are generated, then dashboard shows WA clicks, Maps clicks, and top pages for the selected period."),
        ("US-07", "As an admin, I want to verify data quality.", "Given pending edits exist, then I can approve/reject with notes and maintain source confidence."),
        ("US-08", "As a platform operator, I want paid package controls.", "Given a gym is Starter or Pro, then premium modules are enabled and Free limitations are enforced."),
        ("US-09", "As a search visitor, I want fast mobile pages.", "Given I open a gym page on mobile, then layout loads fast, has no horizontal overflow, and keeps CTA visible."),
    ]
    add_table(doc, ["ID", "User Story", "Acceptance Criteria"], rows, widths=[800, 3800, 4760])


def add_prd_tech_stack(doc):
    add_heading(doc, "10. Tech Stack", 1)
    rows = [
        ("Frontend", "Next.js App Router", "SEO-friendly rendering, route-based pages, strong Vercel support."),
        ("Styling", "Tailwind CSS + design tokens", "Fast iteration, consistent spacing/color/type system."),
        ("Component System", "React components", "Reusable gym cards, CTA blocks, facility tabs, dashboard modules."),
        ("Database", "Postgres via Neon/Supabase", "Relational data fits gym/listing/facility/photo/analytics model."),
        ("ORM", "Prisma or Drizzle", "Typed schema, migrations, safer data access."),
        ("Auth", "Clerk or Supabase Auth", "Owner dashboard and admin access with role separation."),
        ("Storage", "Vercel Blob or S3-compatible storage", "Logo/photo uploads, OG images, gallery assets."),
        ("Search", "Postgres full-text first, Meilisearch later", "MVP can search city/name/facilities without early complexity."),
        ("Analytics", "Vercel Analytics + PostHog", "Track CTA, filters, funnels, and dashboard reporting."),
        ("Payments", "Xendit/Midtrans for Indonesia", "Better local payment support for Starter/Pro packages."),
        ("Deployment", "Vercel", "Fast previews, static/edge delivery, automatic deployments."),
        ("Monitoring", "Vercel logs + uptime monitor", "Catch deploy errors, performance regression, and broken pages."),
    ]
    add_table(doc, ["Layer", "Recommendation", "Justification"], rows, widths=[1800, 3000, 4560])


def add_prd_timeline(doc):
    add_heading(doc, "11. Timeline", 1)
    rows = [
        ("Sprint 1", "Weeks 1-2", "Dataset cleanup, source confidence, PRD/design system", "Milestone 1: verified data model and source rules."),
        ("Sprint 2", "Weeks 3-4", "Directory IA, search/filter prototype, gym card UX", "Clickable public directory prototype."),
        ("Sprint 3", "Weeks 5-6", "Gym detail page, CTA generator, Maps/Instagram integration", "Milestone 2: first 20 public pages preview-ready."),
        ("Sprint 4", "Weeks 7-8", "Dashboard claim flow, gym profile editor", "Claim and edit flow usable internally."),
        ("Sprint 5", "Weeks 9-10", "Photo upload, moderation queue, admin edit review", "Milestone 3: dashboard pilot usable with selected gyms."),
        ("Sprint 6", "Weeks 11-12", "Analytics tracking and dashboard reporting", "CTA/Maps/gallery events visible."),
        ("Sprint 7", "Weeks 13-14", "Starter/Pro package gating and billing workflow", "Milestone 4: monetization pilot ready."),
        ("Sprint 8", "Weeks 15-16", "SEO city/category pages, sitemap, schema, metadata", "Organic landing pages ready for indexing."),
        ("Sprint 9", "Weeks 17-18", "QA, accessibility, Core Web Vitals, browser testing", "Milestone 5: production readiness pass."),
        ("Sprint 10", "Weeks 19-20", "Launch, reporting, outreach kit, post-launch iteration", "MVP public launch and owner outreach."),
    ]
    add_table(doc, ["Sprint", "Duration", "Focus", "Milestone / Output"], rows, widths=[1100, 1200, 3600, 3460])


def add_prd_risks(doc):
    add_heading(doc, "12. Risiko & Mitigasi", 1)
    rows = [
        ("R-01", "Unverified business claims", "5", "4", "20", "Use source confidence, admin moderation, and 'pending verification' labels."),
        ("R-02", "Gym refuses logo/photo usage", "4", "3", "12", "Use text logo fallback and request official media kit during claim process."),
        ("R-03", "Low gym owner adoption", "4", "4", "16", "Start with free verified listings and show lead data before upsell."),
        ("R-04", "SEO takes longer than expected", "4", "4", "16", "Use city/category pages, schema, internal links, and Google Search Console monitoring."),
        ("R-05", "Spam or low-quality leads", "3", "3", "9", "Prefilled message structure and basic anti-spam/rate limiting."),
        ("R-06", "Dashboard data abuse", "4", "2", "8", "Role-based auth, edit moderation, audit logs."),
        ("R-07", "Performance drops from images", "3", "4", "12", "Image compression, lazy loading, explicit dimensions, CDN storage."),
        ("R-08", "Monetization pricing mismatch", "3", "3", "9", "Pilot with flexible monthly pricing and collect owner feedback before scaling."),
    ]
    add_table(doc, ["ID", "Risk", "Impact", "Probability", "Score", "Mitigation"], rows, widths=[700, 2700, 800, 1000, 800, 3360])


def add_gym_specific_appendix(doc, gym):
    add_heading(doc, "Appendix A. Gym-Specific Rollout Notes", 1)
    add_paragraph(doc, f"Gym: {gym['name']}")
    add_paragraph(doc, f"Positioning: {gym['positioning']}")
    add_paragraph(doc, f"Hero headline: {gym['hero']['headline']}")
    add_paragraph(doc, f"Hero subtitle: {gym['hero']['subhead']}")
    add_paragraph(doc, f"Primary CTA: {gym['hero']['cta_primary']} -> {gym['whatsapp']}")
    add_paragraph(doc, f"Secondary CTA: {gym['hero']['cta_secondary']} -> {gym['maps']}")
    add_heading(doc, "Facilities / Programs To Prioritize", 2)
    for item in gym["facility_focus"]:
        add_bullet(doc, f"Facility: {item}")
    for item in gym["program_focus"]:
        add_bullet(doc, f"Program: {item}")
    add_appendix(doc, gym)


def build_doc(gym):
    doc = setup_document()
    add_header_footer(doc, gym)
    add_cover_page(doc, gym)
    add_prd_executive_summary(doc, gym)
    add_prd_background(doc, gym)
    add_prd_okrs(doc)
    add_prd_personas(doc)
    add_prd_scope_roadmap(doc)
    add_prd_functional(doc, gym)
    add_prd_nonfunctional(doc)
    add_prd_model_data(doc)
    add_prd_user_stories(doc)
    add_prd_tech_stack(doc)
    add_prd_timeline(doc)
    add_prd_risks(doc)
    add_gym_specific_appendix(doc, gym)
    return doc


def write_index(generated):
    index = OUT_DIR / "PRD_INDEX.md"
    lines = [
        "# 20 Gym PRD Documents",
        "",
        f"Generated: {DOC_DATE}",
        "",
        "| No | Gym | PRD File |",
        "|---:|---|---|",
    ]
    for gym, path in generated:
        lines.append(f"| {gym['no']} | {gym['name']} | [{path.name}](./{path.name}) |")
    index.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return index


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = []
    for gym in GYMS:
        doc = build_doc(gym)
        filename = f"{gym['no']:02d}-{gym['slug']}-prd.docx"
        path = OUT_DIR / filename
        doc.save(path)
        generated.append((gym, path))
    index = write_index(generated)
    print(f"Generated {len(generated)} DOCX PRDs in {OUT_DIR}")
    print(f"Index: {index}")
    for _, path in generated:
        print(path)


if __name__ == "__main__":
    main()
