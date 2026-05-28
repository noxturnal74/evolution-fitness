from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "product-spec-scroll-animation-v1.0"
DOC_DATE = "May 28, 2026"
PRODUCT = "Gym Growth Platform - Bidirectional Scroll Animation System"
VERSION = "v1.0"
STATUS = "Draft for design, engineering, QA, and implementation"
REVIEWERS = "Albert Saputra, UI/UX Designer, Frontend Engineer, QA Engineer, Performance Reviewer"


def rgb(hex_value):
    hex_value = hex_value.strip("#")
    return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))


BLACK = rgb("#111827")
MUTED = rgb("#4B5563")
BLUE = rgb("#2563EB")
DARK_BLUE = rgb("#1E3A8A")
LIGHT = "EEF2FF"
BORDER = "CBD5E1"


def set_run_font(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    return doc


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def add_cover(doc, doc_type, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(doc_type), size=24, color=BLACK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    set_run_font(p.add_run(subtitle), size=15, color=DARK_BLUE, bold=True)

    rows = [
        ("Product", PRODUCT),
        ("Version", VERSION),
        ("Status", STATUS),
        ("Date", DOC_DATE),
        ("Reviewers", REVIEWERS),
        ("Core Requirement", "Animasi harus tetap aktif saat user scroll ke bawah dan saat user scroll balik ke atas, dengan performa stabil dan respect prefers-reduced-motion."),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        cells = table.rows[i].cells
        set_cell_shading(cells[0], LIGHT)
        for cell in cells:
            set_cell_margins(cell)
        set_run_font(cells[0].paragraphs[0].add_run(label), size=10, color=DARK_BLUE, bold=True)
        set_run_font(cells[1].paragraphs[0].add_run(value), size=10, color=BLACK)
    doc.add_page_break()


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    size = 16 if level == 1 else 13
    color = BLUE if level == 1 else DARK_BLUE
    set_run_font(p.add_run(text), size=size, color=color, bold=True)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    set_run_font(p.add_run(text), size=11, color=BLACK)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text), size=11, color=BLACK)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run(text), size=11, color=BLACK)


def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(t)
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT)
        set_cell_margins(cell)
        set_run_font(cell.paragraphs[0].add_run(header), size=9.5, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            set_cell_margins(cells[idx])
            set_run_font(cells[idx].paragraphs[0].add_run(str(value)), size=9.5, color=BLACK)
    doc.add_paragraph()
    return t


def callout(doc, title, text):
    t = doc.add_table(rows=1, cols=1)
    set_table_borders(t)
    cell = t.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, 140, 140, 180, 180)
    set_run_font(cell.paragraphs[0].add_run(title), size=11, color=DARK_BLUE, bold=True)
    p = cell.add_paragraph()
    set_run_font(p.add_run(text), size=10.5, color=BLACK)
    doc.add_paragraph()


def footer(doc, name):
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(p.add_run(f"{name} | {VERSION}"), size=9, color=MUTED)


def build_prd():
    doc = setup_doc()
    footer(doc, "PRD")
    add_cover(doc, "PRODUCT REQUIREMENTS DOCUMENT", "Bidirectional Scroll Animation For 20 Gym Landing Pages")

    heading(doc, "1. Ringkasan Eksekutif")
    para(doc, "Fitur bidirectional scroll animation memastikan halaman tetap terasa hidup saat user scroll ke bawah maupun scroll balik ke atas. Saat ini banyak animasi website hanya muncul sekali ketika elemen pertama kali terlihat, lalu tidak bereaksi lagi ketika user kembali ke section sebelumnya. Untuk landing page komersial, terutama 20 website gym yang panjang dan visual, pengalaman ini membuat halaman terasa mati setelah user melewati hero.")
    callout(doc, "Product Pitch", "Setiap section harus punya motion yang halus, premium, dan reversible: turun untuk reveal, naik untuk re-activate, tanpa mengganggu performa, aksesibilitas, CTA, atau keterbacaan.")

    heading(doc, "2. Problem Statement")
    bullet(doc, "Animasi one-shot membuat halaman terasa statis setelah user melewati section.")
    bullet(doc, "Scroll balik ke atas sering tidak punya feedback visual, padahal user sedang membandingkan section.")
    bullet(doc, "Animasi berlebihan dapat membuat halaman berat, terutama mobile.")
    bullet(doc, "User dengan reduced motion harus tetap mendapat pengalaman yang nyaman.")

    heading(doc, "3. Goals, Non-Goals, dan Metrics")
    table(doc, ["Category", "Detail", "Target"], [
        ("Goal", "Animasi aktif saat scroll down dan scroll up", "100% section utama punya state enter, leave, re-enter"),
        ("Goal", "Motion terasa premium dan tidak gaming", "Review visual pass"),
        ("Goal", "Performance stabil", "No jank visible, target 60fps feel"),
        ("Goal", "Accessibility aman", "prefers-reduced-motion didukung"),
        ("Non-goal", "Tidak membuat full 3D berat atau library besar", "Vanilla JS + CSS first"),
        ("Metric", "Interaction quality", "No broken CTA, no layout shift, no horizontal overflow"),
    ])

    heading(doc, "4. User Personas")
    table(doc, ["Persona", "Need", "Animation Impact"], [
        ("Pencari gym", "Membandingkan fasilitas, program, harga, lokasi", "Scroll balik harus membantu orientasi ulang section."),
        ("Pemilik gym", "Melihat websitenya terasa profesional", "Motion membuat halaman terlihat hidup dan premium."),
        ("Admin/sales", "Mengarahkan user ke CTA", "Sticky CTA dan section reveal tidak boleh menutup tombol penting."),
        ("User sensitif motion", "Membaca tanpa gangguan", "Reduced motion harus menghilangkan transform berat."),
    ])

    heading(doc, "5. Scope MVP")
    bullet(doc, "Bidirectional reveal untuk section, card, gallery, CTA, dan timeline blocks.")
    bullet(doc, "Direction-aware class: scroll-down, scroll-up, is-visible, was-visible, is-exiting.")
    bullet(doc, "Animasi masuk ulang saat user scroll balik ke section lama.")
    bullet(doc, "Parallax ringan untuk hero dan image surfaces, disable di mobile low-power/reduced-motion.")
    bullet(doc, "Tidak menambah dependency berat kecuali sudah tersedia.")

    heading(doc, "6. Core User Journey")
    numbered(doc, "User membuka halaman dan melihat hero dengan motion ringan.")
    numbered(doc, "User scroll turun; section muncul bertahap dengan reveal dan parallax kecil.")
    numbered(doc, "User klik/lihat fasilitas, gallery, membership, dan CTA.")
    numbered(doc, "User scroll balik ke atas; section yang masuk viewport kembali hidup dengan arah animasi berbeda.")
    numbered(doc, "User tetap bisa klik WhatsApp/Maps tanpa animasi menutup akses.")

    heading(doc, "7. Acceptance Criteria")
    table(doc, ["ID", "Criteria", "Pass Condition"], [
        ("AC-01", "Scroll down animation", "Elemen masuk viewport dengan class is-visible dan motion down preset."),
        ("AC-02", "Scroll up animation", "Elemen yang kembali masuk viewport memakai preset up atau reverse."),
        ("AC-03", "Repeatable", "Animasi dapat terjadi lebih dari sekali, bukan hanya once."),
        ("AC-04", "Reduced motion", "prefers-reduced-motion membuat transform/parallax nonaktif."),
        ("AC-05", "CTA safety", "Floating CTA, nav, modal, form tetap clickable."),
        ("AC-06", "Performance", "Tidak ada layout shift besar dan tidak ada scroll blocking."),
    ])

    heading(doc, "8. Product Requirements Detail")
    para(doc, "Fitur ini bukan sekadar efek visual. Motion harus menjadi sistem navigasi halus yang membantu user memahami posisi mereka di halaman panjang. Pada website gym, user sering bolak-balik antara fasilitas, paket, kelas, lokasi, dan CTA WhatsApp. Karena itu, motion dua arah harus memperjelas konteks tanpa membuat user kehilangan fokus.")
    table(doc, ["Requirement", "Business Rationale", "UX Rationale"], [
        ("Animation repeat on re-entry", "Landing page terasa lebih premium saat dipresentasikan ke pemilik gym.", "User mendapat feedback visual saat kembali ke section yang sudah dilewati."),
        ("Direction-aware motion", "Membedakan pengalaman dari template parallax biasa.", "Gerakan sesuai arah scroll terasa natural dan membantu orientasi."),
        ("CTA lock zone", "Lead WhatsApp tidak boleh turun karena animasi menutup tombol.", "CTA harus tetap terlihat, stabil, dan mudah ditekan."),
        ("Mobile simplified motion", "Mayoritas prospek lokal membuka dari HP.", "Efek ringan menjaga readability dan performa."),
        ("Reduced motion support", "Mengurangi risiko user tidak nyaman.", "Pengalaman tetap accessible untuk semua user."),
    ])

    heading(doc, "9. Success Metrics")
    table(doc, ["Metric", "Definition", "Target MVP", "Measurement Method"], [
        ("Visual QA pass rate", "Jumlah halaman yang lolos review visual tanpa bug motion besar.", "20/20 pages", "Manual checklist desktop + mobile."),
        ("CTA availability", "CTA utama tetap terlihat/clickable saat scroll.", "100% critical sections", "Manual interaction test."),
        ("Animation repeat coverage", "Section utama animate saat scroll down dan up.", ">= 90% animated modules", "Browser visual test."),
        ("Motion performance", "Tidak ada jank jelas saat scroll normal.", "Pass", "Manual scroll test + no long blocking script pattern."),
        ("Accessibility compliance", "Reduced motion dan focus state tersedia.", "Pass", "CSS/JS inspection + manual preference test."),
        ("Content readability", "Teks tidak tertutup overlay/animasi.", "Pass", "Viewport QA 390px, 768px, 1440px."),
    ])

    heading(doc, "10. Rollout Strategy")
    numbered(doc, "Implementasi pertama dilakukan pada satu halaman benchmark agar pola motion, class naming, dan performa terlihat jelas.")
    numbered(doc, "Setelah stabil, sistem dipindahkan ke shared pattern yang bisa dipakai semua website gym.")
    numbered(doc, "Setiap brand gym mendapat tuning preset berbeda: chain gym lebih clean, hardcore lebih bold, women/wellness lebih soft, cafe/lifestyle lebih warm.")
    numbered(doc, "QA dilakukan per kategori brand sebelum seluruh halaman dianggap siap produksi.")
    numbered(doc, "Jika ada bug pada mobile atau reduced motion, fallback harus menampilkan konten secara statis dan tidak menyembunyikan section.")

    heading(doc, "11. Dependencies And Constraints")
    bullet(doc, "Tidak boleh bergantung pada external animation library berat jika vanilla CSS/JS cukup.")
    bullet(doc, "Tidak boleh membuat image, CTA, navbar, atau modal bergeser terlalu agresif.")
    bullet(doc, "Harus kompatibel dengan struktur statis yang sudah dipakai pada project website gym.")
    bullet(doc, "Harus bisa diterapkan ke halaman yang berbeda struktur tanpa crash.")
    bullet(doc, "Dokumentasi implementasi harus cukup jelas agar bisa dipakai ulang untuk project landing page lain.")

    heading(doc, "12. Risks")
    table(doc, ["Risk", "Impact", "Mitigation"], [
        ("Animasi terlalu ramai", "Website terasa gaming, bukan premium.", "Batasi transform, gunakan opacity + small translate, review visual."),
        ("Scroll jank", "User merasa website berat.", "Passive listener, IntersectionObserver, no expensive layout reads."),
        ("CTA tertutup elemen animasi", "Lead turun.", "CTA stable layer, z-index jelas, no animated overlay above CTA."),
        ("Reduced motion tidak jalan", "Accessibility gagal.", "CSS media query + JS motion preference check."),
        ("Terlalu sama antar gym", "Semua terlihat template.", "Preset motion dan rhythm disesuaikan brand category."),
    ])
    return doc


def build_srs():
    doc = setup_doc()
    footer(doc, "SRS")
    add_cover(doc, "SOFTWARE REQUIREMENTS SPECIFICATION", "Bidirectional Scroll Animation Functional And Non-Functional Requirements")

    heading(doc, "1. Purpose")
    para(doc, "SRS ini mendefinisikan kebutuhan software untuk sistem animasi scroll dua arah pada 20 website landing page gym. Fokusnya adalah perilaku, state, input/output, constraint, non-functional requirements, dan acceptance detail.")

    heading(doc, "2. Functional Requirements")
    table(doc, ["ID", "Requirement", "Description", "Priority"], [
        ("FR-01", "Direction detection", "System harus mendeteksi apakah user scroll down atau scroll up berdasarkan posisi scroll terakhir.", "Must"),
        ("FR-02", "Section registry", "Semua elemen reveal harus terdaftar lewat selector data-animate atau class reveal.", "Must"),
        ("FR-03", "Repeated reveal", "Element boleh animate ulang saat re-enter viewport, kecuali ditandai data-animate-once.", "Must"),
        ("FR-04", "Viewport threshold", "Animation trigger harus memakai IntersectionObserver dengan threshold dan rootMargin yang bisa disetel.", "Must"),
        ("FR-05", "Exit state", "Saat elemen keluar viewport, system harus memberi class is-exiting atau menghapus is-visible sesuai preset.", "Should"),
        ("FR-06", "Preset support", "Support fade-up, fade-down, slide-left, slide-right, scale-in, parallax-soft, card-stagger.", "Should"),
        ("FR-07", "Stagger", "Card group harus bisa muncul bertahap tanpa delay berlebihan.", "Should"),
        ("FR-08", "Analytics hook", "System menyediakan custom event animation:enter dan animation:reenter untuk debug/analytics.", "Could"),
    ])

    heading(doc, "3. Non-Functional Requirements")
    table(doc, ["Area", "Requirement", "Target"], [
        ("Performance", "Scroll handler tidak boleh melakukan layout thrashing.", "Use requestAnimationFrame and passive listeners."),
        ("Accessibility", "Respect prefers-reduced-motion.", "Motion becomes opacity-only or instant."),
        ("Compatibility", "Works on modern Chrome, Edge, Safari, Firefox mobile/desktop.", "Progressive enhancement."),
        ("Maintainability", "Animation presets terpusat di CSS tokens/classes.", "No hard-coded inline animation per section."),
        ("Reliability", "No console error if elements missing.", "Defensive selectors."),
        ("Responsiveness", "Mobile animation lebih ringan.", "No heavy parallax under 768px."),
    ])

    heading(doc, "4. External Interface Requirements")
    bullet(doc, "HTML uses data-animate, data-animate-group, data-animate-once, and data-parallax-speed where needed.")
    bullet(doc, "CSS exposes animation classes and reduced motion fallback.")
    bullet(doc, "JS exposes window.ScrollMotionController for debug only, not required for normal use.")
    bullet(doc, "No external API needed.")

    heading(doc, "5. Data Requirements")
    table(doc, ["Data", "Type", "Usage"], [
        ("lastScrollY", "number", "Compare direction."),
        ("scrollDirection", "up/down", "Apply direction-aware classes."),
        ("visibleElements", "Set<Element>", "Track elements currently in viewport."),
        ("animationPreset", "string", "Map element to CSS preset."),
        ("motionPreference", "normal/reduced", "Disable heavy effects."),
    ])

    heading(doc, "6. Acceptance Tests")
    numbered(doc, "Scroll from hero to footer: all reveal sections animate smoothly.")
    numbered(doc, "Scroll from footer back to hero: prior sections animate or reverse as they re-enter.")
    numbered(doc, "Enable reduced motion in OS/browser: no parallax/large transform is active.")
    numbered(doc, "Open gallery/modal while scrolled mid-page: animation does not break focus or close control.")
    numbered(doc, "Resize mobile viewport: no horizontal overflow and animation remains light.")

    heading(doc, "7. Detailed Requirement Rules")
    table(doc, ["Rule ID", "Rule", "Required Behavior", "Failure Condition"], [
        ("R-01", "JS enhancement", "Content must be readable before JS initializes.", "Section remains invisible if JS fails."),
        ("R-02", "Class lifecycle", "Element can move from reveal -> is-visible -> is-exiting -> is-visible again.", "Element only animates once without opt-in."),
        ("R-03", "Direction tolerance", "Small scroll jitter under 4px should not flip direction.", "Animation flickers when touch scroll bounces."),
        ("R-04", "Viewport margin", "Animation should start slightly before section center.", "Content appears too late after user has already read empty space."),
        ("R-05", "Touch devices", "Touch scroll must not be blocked.", "Listener prevents default scrolling."),
        ("R-06", "Modal coexistence", "Opening modal should pause background interaction, not break observer.", "Focus trapped behind overlay or Escape fails."),
    ])

    heading(doc, "8. Error Handling")
    bullet(doc, "If IntersectionObserver is unavailable, all reveal elements become visible immediately.")
    bullet(doc, "If an element has unknown data-animate preset, default to fade-up.")
    bullet(doc, "If CSS variables are missing, fallback timing and transform values must still work.")
    bullet(doc, "If reduced motion is active after page load, controller should respond to media query change.")
    bullet(doc, "If a section is dynamically inserted later, init() can be called again without duplicating observers.")

    heading(doc, "9. Browser And Device Matrix")
    table(doc, ["Viewport/Browser", "Expected Behavior", "Priority"], [
        ("Chrome desktop 1440px", "Full bidirectional reveal and light parallax.", "P0"),
        ("Edge desktop 1366px", "Same as Chrome, no layout overflow.", "P0"),
        ("Safari iOS 390px", "Simplified reveal, no heavy fixed background parallax.", "P0"),
        ("Chrome Android 412px", "Touch scroll smooth, CTA fixed/floating accessible.", "P0"),
        ("Firefox desktop", "Reveal works; parallax fallback acceptable.", "P1"),
        ("Reduced motion browser setting", "No transform-heavy animation.", "P0"),
    ])

    heading(doc, "10. Traceability Matrix")
    table(doc, ["Business Need", "Functional Requirement", "Acceptance Test"], [
        ("Website terasa hidup saat scroll balik", "FR-03 Repeated reveal", "Acceptance Test 2"),
        ("Tidak membuat mobile berat", "NFR Performance + mobile simplified motion", "Acceptance Test 5"),
        ("Tetap accessible", "Reduced motion support", "Acceptance Test 3"),
        ("CTA tetap menjual", "CTA safety and no overlay blocking", "Acceptance Test 4"),
        ("Bisa dipakai semua gym", "Section registry and preset support", "Batch validation test"),
    ])
    return doc


def build_sdd():
    doc = setup_doc()
    footer(doc, "SDD")
    add_cover(doc, "SOFTWARE DESIGN DOCUMENT", "Architecture And Implementation Design For Scroll Animation System")

    heading(doc, "1. Architecture Overview")
    para(doc, "Sistem dirancang sebagai enhancement ringan di atas HTML/CSS/JS statis. Jika JavaScript gagal, konten tetap terlihat. Jika JavaScript aktif, body mendapat class js-ready lalu controller mengatur class animasi berdasarkan arah scroll dan visibility.")

    heading(doc, "2. Components")
    table(doc, ["Component", "Responsibility", "File Target"], [
        ("ScrollMotionController", "Main orchestrator untuk observer, scroll direction, class lifecycle.", "script.js"),
        ("DirectionTracker", "Menyimpan lastScrollY dan menghitung up/down dengan tolerance.", "script.js"),
        ("SectionObserver", "IntersectionObserver untuk reveal/re-enter.", "script.js"),
        ("MotionPresetRegistry", "Menerjemahkan data-animate menjadi CSS class.", "styles.css + script.js"),
        ("AccessibilityAdapter", "Cek prefers-reduced-motion dan disable heavy transform.", "script.js"),
        ("PerformanceGuard", "Passive scroll listener, requestAnimationFrame, no layout-heavy reads.", "script.js"),
    ])

    heading(doc, "3. State Model")
    table(doc, ["State", "Meaning", "Class"], [
        ("Initial", "Element belum pernah masuk viewport.", "reveal"),
        ("EnteringDown", "Element masuk viewport saat scroll turun.", "is-visible scroll-down"),
        ("EnteringUp", "Element masuk viewport saat scroll naik.", "is-visible scroll-up"),
        ("Exiting", "Element keluar viewport.", "is-exiting"),
        ("ReducedMotion", "Motion preference reduced.", "motion-reduced"),
    ])

    heading(doc, "4. Algorithm")
    numbered(doc, "On load, add body class js-ready.")
    numbered(doc, "Collect animated elements using [data-animate], .reveal, .bento-card, .gallery-item, and major modules.")
    numbered(doc, "Attach passive scroll listener and update scrollDirection with requestAnimationFrame.")
    numbered(doc, "IntersectionObserver calls enter/exit handlers.")
    numbered(doc, "On enter, remove exit classes, set direction class, set is-visible.")
    numbered(doc, "On exit, either remove is-visible for repeatable elements or preserve if data-animate-once.")
    numbered(doc, "If reduced motion, skip transform classes and only set visible state.")

    heading(doc, "5. Pseudocode")
    callout(doc, "Core Flow", "lastY = scrollY; onScroll -> direction = scrollY > lastY ? down : up; observer.onEnter(el) -> el.classList.add('is-visible', 'scroll-' + direction); observer.onExit(el) -> if repeatable remove is-visible and add is-exiting.")

    heading(doc, "6. CSS Design")
    bullet(doc, "Default .reveal is visible-safe when JS is off.")
    bullet(doc, ".js-ready .reveal starts hidden/offset.")
    bullet(doc, ".reveal.is-visible.scroll-down uses translateY(0) from positive offset.")
    bullet(doc, ".reveal.is-visible.scroll-up uses translateY(0) from negative offset or subtle reverse preset.")
    bullet(doc, "@media (prefers-reduced-motion: reduce) disables transition, transform, parallax.")

    heading(doc, "7. File-Level Design")
    table(doc, ["File", "Responsibility", "Notes"], [
        ("index.html", "Add data-animate attributes to major modules.", "Markup remains semantic and readable without JS."),
        ("styles.css", "Define motion tokens, presets, reduced-motion, responsive tuning.", "No repeated per-section hard-coded transitions."),
        ("script.js", "Initialize controller and interaction modules.", "Defensive DOM access and no console spam."),
        ("build-check.mjs", "Validate required hooks and reduced motion.", "Optional but recommended for production gates."),
        ("README.md", "Document available presets and usage examples.", "Helps future websites reuse system."),
    ])

    heading(doc, "8. CSS Token Proposal")
    table(doc, ["Token", "Default", "Usage"], [
        ("--motion-duration-fast", "260ms", "Small labels, badges, micro UI."),
        ("--motion-duration-base", "520ms", "Section headings and cards."),
        ("--motion-duration-slow", "760ms", "Hero and large visual modules."),
        ("--motion-ease", "cubic-bezier(.2,.8,.2,1)", "Premium responsive easing."),
        ("--motion-distance", "32px", "Default reveal offset."),
        ("--motion-distance-mobile", "18px", "Mobile reveal offset."),
    ])

    heading(doc, "9. Integration Pattern")
    callout(doc, "HTML Example", "<section data-animate='fade-up'>...</section> and <div data-animate-group='stagger-cards'>...</div> are enough for normal use. Advanced sections can add data-parallax-speed='0.12'.")
    callout(doc, "JS Example", "ScrollMotionController.init({ selector: '[data-animate], .reveal', repeat: true, threshold: 0.16, rootMargin: '0px 0px -10% 0px' })")

    heading(doc, "10. Performance Design")
    bullet(doc, "Use IntersectionObserver for visibility instead of checking all elements on every scroll.")
    bullet(doc, "Use one passive scroll listener only to track direction.")
    bullet(doc, "Batch DOM writes inside requestAnimationFrame.")
    bullet(doc, "Avoid reading getBoundingClientRect repeatedly during active scroll except for controlled parallax elements.")
    bullet(doc, "Disable fixed background attachment and heavy parallax on mobile.")

    heading(doc, "11. Security And Reliability Notes")
    bullet(doc, "No user input is executed as code.")
    bullet(doc, "Generated WhatsApp text remains URL-encoded by existing CTA generator.")
    bullet(doc, "Controller must not fetch remote scripts or depend on third-party animation CDN.")
    bullet(doc, "Unknown elements should be ignored safely.")

    heading(doc, "12. QA Hooks")
    bullet(doc, "Add temporary data-motion-debug attribute support only for development if needed.")
    bullet(doc, "Expose a lightweight window.ScrollMotionController.getState() for debugging when available.")
    bullet(doc, "Build checks should search for js-ready, IntersectionObserver, Escape key modal handling, and prefers-reduced-motion.")
    return doc


def build_uiux():
    doc = setup_doc()
    footer(doc, "UI UX FLOW")
    add_cover(doc, "UI/UX FLOW DOCUMENT", "Scroll Up And Down Animation Experience Flow")

    heading(doc, "1. UX Principle")
    para(doc, "Motion harus membantu orientasi, bukan menjadi dekorasi berlebihan. User harus merasa halaman responsif saat bergerak ke bawah dan saat kembali ke atas. Setiap animasi harus mendukung scanning, comparison, dan CTA.")

    heading(doc, "2. Primary Flow")
    table(doc, ["Step", "User Action", "UI Response", "UX Goal"], [
        ("1", "Landing di hero", "Hero media, title, CTA muncul dengan calm cinematic reveal.", "Bangun first impression."),
        ("2", "Scroll ke bawah", "Section baru masuk dengan fade/slide/stagger.", "Arahkan mata ke konten berikutnya."),
        ("3", "Hover/click gambar", "Caption atau lightbox muncul tanpa menutup CTA utama.", "Gambar menjadi informatif."),
        ("4", "Scroll balik ke atas", "Section re-enter dengan reverse/up animation.", "Membantu user membandingkan ulang."),
        ("5", "Klik CTA", "CTA tetap stabil dan fokus tidak terganggu motion.", "Konversi tetap jelas."),
    ])

    heading(doc, "3. Screen-Level Flow")
    numbered(doc, "Hero: light parallax, CTA stable, status card readable.")
    numbered(doc, "Platform/benchmark sections: cards stagger in both directions.")
    numbered(doc, "Facilities: bento cards reveal with subtle scale; image hover has paragraph detail.")
    numbered(doc, "Plans: package cards animate in but active state remains persistent.")
    numbered(doc, "Gallery: images reveal, lightbox opens, Escape closes.")
    numbered(doc, "Join: form appears after section reveal; keyboard focus remains clear.")

    heading(doc, "4. Motion Preset Matrix")
    table(doc, ["Element", "Scroll Down", "Scroll Up", "Reduced Motion"], [
        ("Hero media", "Parallax shift downward", "Parallax returns upward", "Static image"),
        ("Section heading", "Fade up", "Fade down / reverse fade", "Instant visible"),
        ("Cards", "Stagger fade up", "Stagger fade down", "No stagger"),
        ("Gallery image", "Scale-in + fade", "Scale-in subtle reverse", "Opacity only"),
        ("CTA", "No movement after visible", "No movement after visible", "Always stable"),
    ])

    heading(doc, "5. Accessibility Flow")
    bullet(doc, "Keyboard user should not lose focus when animation classes change.")
    bullet(doc, "Screen reader content order remains DOM order; animation must not reorder content.")
    bullet(doc, "Reduced motion removes transform and parallax.")
    bullet(doc, "No essential information appears only during hover.")

    heading(doc, "6. Brand-Specific Motion Direction")
    table(doc, ["Gym Type", "Motion Feel", "Visual Treatment"], [
        ("Premium/luxury", "Slow fade, small parallax, calm scale.", "Large imagery, wider spacing, subtle glow."),
        ("Hardcore/bodybuilding", "Sharper slide, stronger contrast, faster reveal.", "Industrial image crop, bold type, dark panels."),
        ("Women/wellness", "Soft fade, gentle vertical movement.", "Warm surfaces, privacy-focused copy, less aggressive motion."),
        ("Cafe/lifestyle", "Smooth editorial reveal, image hover captions.", "Earthy palette, community photos, comfortable rhythm."),
        ("Chain/modern", "Clean bento stagger, app-like micro interaction.", "Facility grid, membership comparator, strong CTA."),
    ])

    heading(doc, "7. Responsive UX Rules")
    bullet(doc, "Desktop can use subtle parallax and larger stagger because screen space supports scanning.")
    bullet(doc, "Tablet should reduce stagger delay so user does not wait for content.")
    bullet(doc, "Mobile should keep animation distance short and avoid fixed background parallax.")
    bullet(doc, "Hero title and CTA must not be animated in a way that causes horizontal overflow.")
    bullet(doc, "Sticky WhatsApp CTA should not cover form fields, final CTA, or footer links.")

    heading(doc, "8. Image Interaction Flow")
    numbered(doc, "User sees image card without paragraph clutter.")
    numbered(doc, "On hover/focus, a short paragraph description slides/fades in inside a readable overlay.")
    numbered(doc, "On click/tap, gallery lightbox opens with the same paragraph available below image title.")
    numbered(doc, "Escape key or close button returns user to the same card.")
    numbered(doc, "On mobile, description appears on tap or is shown as compact caption below the image, not only hover.")

    heading(doc, "9. Copywriting And Motion Pairing")
    table(doc, ["Content Type", "Copy Rule", "Motion Rule"], [
        ("Hero", "Short, confident, brand-specific.", "One controlled entrance, no repeated title jump while in viewport."),
        ("Facilities", "Explain what user can do there.", "Bento cards stagger lightly."),
        ("Membership", "Clear package names and CTA.", "Selected package stays stable; no card jumping."),
        ("Classes", "Use day/time/category if available.", "Filter changes should animate list gently."),
        ("Location", "Use address/source links, no fake claims.", "Map block fades in; buttons remain static."),
    ])

    heading(doc, "10. UX QA Checklist")
    bullet(doc, "Can user understand the page in 5 seconds at hero?")
    bullet(doc, "Can user reach WhatsApp from hero, membership, and final CTA?")
    bullet(doc, "Does scrolling back up feel intentional instead of dead/static?")
    bullet(doc, "Does every hover image paragraph remain readable and not cover critical content?")
    bullet(doc, "Does mobile feel lighter than desktop while preserving premium feel?")
    return doc


def build_tasks():
    doc = setup_doc()
    footer(doc, "TASK BREAKDOWN")
    add_cover(doc, "TASK BREAKDOWN", "Implementation Plan For Bidirectional Scroll Animation")

    heading(doc, "1. Epic Breakdown")
    table(doc, ["Epic", "Description", "Owner"], [
        ("E1 Motion System", "Build reusable scroll direction and observer controller.", "Frontend"),
        ("E2 CSS Presets", "Create animation classes, reduced motion, responsive variants.", "Frontend + Design"),
        ("E3 Site Integration", "Apply to 20 gym pages and key modules.", "Frontend"),
        ("E4 QA & Performance", "Validate no jank, no broken CTA, no mobile overflow.", "QA"),
        ("E5 Documentation", "Document usage, do/don't, and acceptance checks.", "Product"),
    ])

    heading(doc, "2. Task List")
    table(doc, ["ID", "Task", "Acceptance", "Estimate"], [
        ("T-01", "Audit current reveal/parallax code", "List current selectors and one-shot behavior.", "0.5d"),
        ("T-02", "Create ScrollMotionController", "Direction, observer, enter/exit state working.", "1d"),
        ("T-03", "Add CSS bidirectional presets", "scroll-up and scroll-down classes visually distinct but subtle.", "1d"),
        ("T-04", "Add reduced motion fallback", "All transform/parallax disabled under prefers-reduced-motion.", "0.5d"),
        ("T-05", "Integrate with modules", "Hero, platform, benchmark, facility, gallery, plans, join animate.", "1d"),
        ("T-06", "Mobile tuning", "No horizontal overflow; motion lighter under 768px.", "0.5d"),
        ("T-07", "Interaction regression", "Menus, modals, tabs, forms, WhatsApp generator still work.", "1d"),
        ("T-08", "20-site batch validation", "All build-check scripts pass.", "0.5d"),
        ("T-09", "Visual QA", "Check at least representative sites: chain, hardcore, wellness, cafe, local.", "1d"),
        ("T-10", "Documentation update", "README/report explains motion system and how to tune presets.", "0.5d"),
    ])

    heading(doc, "3. Sprint Plan")
    table(doc, ["Sprint", "Focus", "Deliverable"], [
        ("Sprint 1", "Core motion engine and CSS tokens", "Controller + presets in one site."),
        ("Sprint 2", "Batch integration", "20 pages updated and build passed."),
        ("Sprint 3", "QA and polish", "Visual review, mobile tuning, accessibility pass."),
    ])

    heading(doc, "4. Definition Of Done")
    bullet(doc, "All 20 sites pass build/validation.")
    bullet(doc, "Scroll down and scroll up both trigger visible animation changes.")
    bullet(doc, "Reduced motion is respected.")
    bullet(doc, "No broken buttons, modals, gallery, or WhatsApp generator.")
    bullet(doc, "No horizontal overflow on mobile.")
    bullet(doc, "Implementation notes are documented.")

    heading(doc, "5. Detailed 10-Sprint Breakdown")
    table(doc, ["Sprint", "Duration", "Goal", "Output"], [
        ("1", "1 week", "Audit current animation and interaction code.", "Motion audit matrix and bug list."),
        ("2", "1 week", "Create shared motion tokens and CSS presets.", "styles.css motion layer."),
        ("3", "2 weeks", "Build ScrollMotionController.", "Direction-aware observer and lifecycle classes."),
        ("4", "2 weeks", "Integrate benchmark site.", "One fully working representative landing page."),
        ("5", "2 weeks", "Apply to chain/modern gyms.", "FTL, Fitness Plus, FITX, Osbond tuned."),
        ("6", "2 weeks", "Apply to local/community gyms.", "Evolution, Planet, New Icon, M Gym, DM Gym tuned."),
        ("7", "2 weeks", "Apply to hardcore/performance gyms.", "Draco, SpeedRocky, Warriors, BlackBox, Champion tuned."),
        ("8", "2 weeks", "Apply to wellness/cafe/women-focused gyms.", "Belle Crown, Crystal, OCIGEN, Optimum, Glanzfit tuned."),
        ("9", "3 weeks", "QA, accessibility, mobile, performance.", "Bug fixes and validation report."),
        ("10", "3 weeks", "Polish, documentation, deployment prep.", "Production-ready batch and handover notes."),
    ])

    heading(doc, "6. RACI")
    table(doc, ["Workstream", "Responsible", "Accountable", "Consulted", "Informed"], [
        ("Motion system", "Frontend Engineer", "Tech Lead", "UI Designer, QA", "Product Owner"),
        ("Visual tuning", "UI Designer", "Creative Director", "Brand Strategist", "Frontend Engineer"),
        ("Copy and data", "Brand Strategist", "Product Owner", "Researcher", "Sales/Admin"),
        ("Accessibility", "QA Engineer", "Tech Lead", "UI Designer", "Product Owner"),
        ("Deployment", "Frontend Engineer", "Tech Lead", "QA", "Product Owner"),
    ])

    heading(doc, "7. QA Task Checklist")
    bullet(doc, "Check scroll down reveal on hero, facilities, membership, gallery, location, FAQ, final CTA.")
    bullet(doc, "Check scroll up re-entry from footer back to hero.")
    bullet(doc, "Check reduced motion with browser/OS setting.")
    bullet(doc, "Check keyboard: nav, modal, FAQ, gallery, package selector, CTA.")
    bullet(doc, "Check mobile 390px and desktop 1440px.")
    bullet(doc, "Check all buttons: WhatsApp, Maps, Instagram, package selector, class/facility tabs.")
    bullet(doc, "Check image hover/tap paragraph descriptions.")
    bullet(doc, "Check no broken images, no fallback logo unless source missing.")

    heading(doc, "8. Implementation Deliverables")
    table(doc, ["Deliverable", "Format", "Owner", "Done When"], [
        ("Motion system code", "script.js + styles.css", "Frontend", "Reusable and documented."),
        ("Updated landing pages", "20 static sites", "Frontend", "All sections integrated."),
        ("Validation script", "build-check.mjs", "Frontend/QA", "Fails on missing essentials."),
        ("QA report", "Markdown", "QA", "Lists pass/fail per gym."),
        ("Handover docs", "README + specs", "Product", "Non-technical user can understand next steps."),
    ])
    return doc


DOC_BUILDERS = [
    ("01-PRD-bidirectional-scroll-animation.docx", build_prd),
    ("02-SRS-bidirectional-scroll-animation.docx", build_srs),
    ("03-SDD-bidirectional-scroll-animation.docx", build_sdd),
    ("04-UI-UX-FLOW-bidirectional-scroll-animation.docx", build_uiux),
    ("05-TASK-BREAKDOWN-bidirectional-scroll-animation.docx", build_tasks),
]


def write_index(paths):
    lines = [
        "# Bidirectional Scroll Animation Spec Package",
        "",
        f"Product: {PRODUCT}",
        f"Version: {VERSION}",
        f"Date: {DOC_DATE}",
        "",
        "| No | Document | File |",
        "|---:|---|---|",
    ]
    names = ["PRD", "SRS", "SDD", "UI/UX Flow", "Task Breakdown"]
    for idx, path in enumerate(paths, start=1):
        lines.append(f"| {idx} | {names[idx - 1]} | [{path.name}](./{path.name}) |")
    (OUT_DIR / "README.md").write_text("\n".join(lines) + "\n", encoding="utf-8")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    generated = []
    for filename, builder in DOC_BUILDERS:
        doc = builder()
        path = OUT_DIR / filename
        doc.save(path)
        generated.append(path)
    write_index(generated)
    print(f"Generated {len(generated)} documents in {OUT_DIR}")
    for path in generated:
        print(path)


if __name__ == "__main__":
    main()
